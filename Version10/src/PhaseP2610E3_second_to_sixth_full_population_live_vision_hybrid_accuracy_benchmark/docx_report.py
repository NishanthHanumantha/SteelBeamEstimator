"""Stakeholder DOCX from report_data.json. Hybrid architecture content. Does not overwrite QA.30."""
from __future__ import annotations

from pathlib import Path
from typing import Any, Dict, List, Optional

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

from .config import DOCX_NAME, GATE_VERSION, INCLUDED_SET_KEYS, MODEL_VERSION, PHASE_ID

NAVY = RGBColor(0x1B, 0x36, 0x5D)
STEEL = RGBColor(0x2F, 0x5D, 0x7C)
RED = RGBColor(0x9B, 0x2C, 0x2C)
AMBER = RGBColor(0x8A, 0x5A, 0x00)
SLATE = RGBColor(0x3D, 0x4A, 0x57)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
ROW_ALT = "F4F7FA"
HEADER_FILL = "1B365D"
ACCENT_FILL = "2F5D7C"
GREEN_FILL = "1F6B3A"
GOLD_FILL = "B88A2E"
RED_FILL = "9B2C2C"


def _set_run(run, *, size=11, bold=False, color=SLATE, font="Calibri"):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = color


def _shade(cell, hex_color: str) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def _borders(cell, color="C5CDD6") -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        tcBorders.append(el)
    tcPr.append(tcBorders)


def cell_text(cell, text, *, bold=False, color=SLATE, size=9, align="center", fill=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
    }[align]
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run(str(text))
    _set_run(run, size=size, bold=bold, color=color)
    if fill:
        _shade(cell, fill)
    _borders(cell)


def set_widths(table, widths_cm):
    table.autofit = False
    table.allow_autofit = False
    total = sum(widths_cm)
    tblPr = table._tbl.tblPr if table._tbl.tblPr is not None else OxmlElement("w:tblPr")
    tblW = tblPr.find(qn("w:tblW"))
    if tblW is None:
        tblW = OxmlElement("w:tblW")
        tblPr.append(tblW)
    tblW.set(qn("w:w"), str(int(total * 567)))
    tblW.set(qn("w:type"), "dxa")
    for row in table.rows:
        for i, w in enumerate(widths_cm):
            if i < len(row.cells):
                row.cells[i].width = Cm(w)


def heading_bar(doc, title: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell_text(table.cell(0, 0), title, bold=True, color=WHITE, size=11, fill=HEADER_FILL)
    set_widths(table, [17.0])
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(3)


def grid(doc, headers, data, col_widths, *, total_last=False):
    table = doc.add_table(rows=1 + len(data), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(headers):
        cell_text(table.cell(0, j), h, bold=True, color=WHITE, size=7.5, fill=HEADER_FILL)
    for i, row in enumerate(data):
        is_total = total_last and i == len(data) - 1
        fill = "E8EEF4" if is_total else (ROW_ALT if i % 2 else "FFFFFF")
        for j, val in enumerate(row):
            cell_text(
                table.cell(i + 1, j),
                val,
                bold=is_total or j == 0,
                color=NAVY if j == 0 else SLATE,
                size=8,
                align="left" if j == 0 else "center",
                fill=fill,
            )
    set_widths(table, col_widths)
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(6)


def body(doc, text, *, size=10, color=SLATE, space_after=6, bold=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing = 1.08
    run = p.add_run(text)
    _set_run(run, size=size, color=color, bold=bold)
    return p


def kpis(doc, items):
    table = doc.add_table(rows=2, cols=len(items))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    fills = [HEADER_FILL, ACCENT_FILL, GREEN_FILL, GOLD_FILL]
    for i, (label, value) in enumerate(items):
        cell_text(table.cell(0, i), label, bold=True, color=WHITE, size=7.5, fill=fills[i % len(fills)])
        cell_text(table.cell(1, i), value, bold=True, color=NAVY, size=12, fill="FFFFFF")
    set_widths(table, [17.0 / len(items)] * len(items))
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(8)


def _fmt(v: Any, n: int = 2, suffix: str = "") -> str:
    if v is None:
        return "—"
    try:
        return f"{float(v):.{n}f}{suffix}"
    except (TypeError, ValueError):
        return str(v)


def _fmt_int(v: Any) -> str:
    if v is None:
        return "—"
    try:
        return f"{int(v):,}"
    except (TypeError, ValueError):
        return str(v)


def _fmt_kg(v: Any) -> str:
    if v is None:
        return "—"
    try:
        return f"{float(v):,.0f}"
    except (TypeError, ValueError):
        return str(v)


def _fmt_signed_kg(v: Any) -> str:
    if v is None:
        return "—"
    try:
        n = float(v)
        sign = "−" if n < 0 else ("+" if n > 0 else "")
        return f"{sign}{abs(n):,.0f}"
    except (TypeError, ValueError):
        return str(v)


def _diameter_ident_rows(data: Dict[str, Any]) -> List[List[str]]:
    rows = []
    for row in (data.get("diameter_wise") or {}).get("identification_rows") or []:
        rows.append(
            [
                str(row.get("diameter_label") or ""),
                _fmt_int(row.get("gt_bar_lines")),
                _fmt_int(row.get("detected")),
                _fmt_int(row.get("match")),
                _fmt_int(row.get("wrong_diameter")),
                _fmt(row.get("diameter_identification_percent"), suffix="%"),
                str(row.get("note") or ""),
            ]
        )
    return rows


def _diameter_steel_rows(data: Dict[str, Any]) -> List[List[str]]:
    rows = []
    for row in (data.get("diameter_wise") or {}).get("steel_rows") or []:
        ratio = row.get("quantity_ratio_percent")
        rows.append(
            [
                str(row.get("diameter_label") or ""),
                _fmt_kg(row.get("benchmark_kg") if row.get("benchmark_kg") is not None else row.get("estimator_kg")),
                _fmt_kg(row.get("model_kg")),
                _fmt_signed_kg(row.get("difference_kg")),
                _fmt(row.get("difference_pct"), n=1, suffix="%"),
                "—" if ratio is None else f"{float(ratio):.0f}%",
            ]
        )
    return rows


def _add_picture(doc, path: Optional[str], width_in: float = 6.4) -> None:
    if not path or not Path(path).exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Inches(width_in))
    p.paragraph_format.space_after = Pt(8)


def write_docx(*, out_root: Path, data: Dict[str, Any], charts: Dict[str, str]) -> Path:
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Cm(21.0)
    sec.page_height = Cm(29.7)
    sec.left_margin = Cm(1.7)
    sec.right_margin = Cm(1.7)
    sec.top_margin = Cm(1.3)
    sec.bottom_margin = Cm(1.5)

    hp = sec.header.paragraphs[0]
    r = hp.add_run("SteelBeam Estimator  ·  Version 10  ·  Current hybrid architecture")
    _set_run(r, size=8, bold=True, color=STEEL)
    r2 = hp.add_run("     Second–Sixth Sets  ·  First Set excluded  ·  Shadow benchmark")
    _set_run(r2, size=8, color=AMBER)

    fp = sec.footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = fp.add_run(
        "Confidential / internal use  ·  Vision semantics are diagnostic/shadow inputs  ·  Not a production promotion"
    )
    _set_run(fr, size=7.5, color=STEEL)

    pooled = data.get("pooled") or {}
    per_set = data.get("per_set") or {}
    cov = data.get("vision_coverage") or {}
    pop = data.get("population") or {}
    vis_exec = data.get("vision_execution") or {}

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    _set_run(p.add_run("STEEL BEAM ESTIMATION"), size=20, bold=True, color=NAVY)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    _set_run(p.add_run("Current Hybrid Architecture Performance  ·  Second to Sixth Drawing Sets"), size=12, color=STEEL)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    _set_run(
        p.add_run(
            f"Version 10  ·  MODEL_VERSION {MODEL_VERSION}  ·  Gate {GATE_VERSION}  ·  "
            f"Decision {data.get('decision') or '—'}  ·  Benchmark mode {data.get('mode') or '—'}"
        ),
        size=9,
        color=AMBER,
        bold=True,
    )

    heading_bar(doc, "CURRENT HYBRID ARCHITECTURE PERFORMANCE  —  SECOND TO SIXTH SETS")
    kpis(
        doc,
        [
            ("Beam identification", _fmt(pooled.get("beam_identification_percent"), suffix="%")),
            ("Bar identification", _fmt(pooled.get("bar_identification_percent"), suffix="%")),
            ("Correct of detected bars", _fmt(pooled.get("correct_of_detected_percent"), suffix="%")),
            ("Diameter identification", _fmt(pooled.get("diameter_identification_percent"), suffix="%")),
        ],
    )
    kpis(
        doc,
        [
            ("Steel / weight accuracy", _fmt(pooled.get("weight_accuracy_percent"), suffix="%")),
            ("Overall accuracy", _fmt(pooled.get("overall_accuracy_percent"), suffix="%")),
            ("HYBRID coverage", _fmt(cov.get("hybrid_percent"), suffix="%")),
            ("FALLBACK coverage", _fmt(cov.get("fallback_percent"), suffix="%")),
        ],
    )
    body(
        doc,
        "This report measures the current hybrid architecture on Second through Sixth drawing sets. "
        "Claude Vision is the preferred semantic interpreter after validation. The deterministic engine "
        "remains the engineering calculation authority. First Set is excluded. This is not a historical "
        "improvement study and not a production-readiness claim.",
        size=10,
        space_after=8,
    )

    heading_bar(doc, "1.  Executive summary")
    body(
        doc,
        f"Five drawing sets were discovered dynamically and scored with the QA.2A / QA.3.0 methodology. "
        f"Model beams: {pop.get('model_beam_total')}. Ground-truth beams: {pop.get('estimator_beam_total')}. "
        f"Matched IDs: {pop.get('matched_total')}. Vision usable (HYBRID): {cov.get('hybrid_count')}. "
        f"Deterministic fallback: {cov.get('fallback_count')}. "
        f"Pooled overall accuracy is the mean of pooled beam identification, bar identification, "
        f"correct-of-detected, and kg-pooled steel/weight accuracy. Diameter is reported separately.",
        size=10,
    )

    heading_bar(doc, "2.  Current hybrid architecture")
    body(
        doc,
        "DXF → Beam discovery → Deterministic geometry context → Claude Vision semantic interpretation → "
        "D.1 semantic authority contract → D.2 hybrid resolution → D.3 engineering binding → "
        "D.4 shadow engineering calculation → accuracy benchmark versus estimator truth.",
        size=10,
        bold=True,
        space_after=4,
    )
    body(
        doc,
        "VISION SEMANTICS (preferred after validation): target identity, layer, physical group detection, "
        "bar count, diameter, specification, MAIN/EXTRA role, support scope, stirrup identification. "
        "DETERMINISTIC ENGINEERING: spacers, geometry, cut length, development length, anchorage, hooks/bends, "
        "stirrup geometry/zones/spacing/quantity, piece generation, weight, BBS, workbook generation. "
        "Vision-preferred does not mean blindly accepted. Ambiguous matches are withheld, not forced.",
        size=9.5,
    )

    heading_bar(doc, "3.  Benchmark scope and population")
    pop_rows = []
    for key in INCLUDED_SET_KEYS:
        row = (pop.get("by_set") or {}).get(key) or per_set.get(key) or {}
        pop_rows.append(
            [
                f"{key} Set",
                str(row.get("model_beams") or row.get("discovered_model_beam_count") or "—"),
                str(row.get("gt_beams") or row.get("discovered_estimator_beam_count") or "—"),
                str(row.get("matched_beams") or row.get("matched_benchmark_population") or "—"),
                str(len(row.get("unmatched_model_beams") or []) if isinstance(row.get("unmatched_model_beams"), list) else row.get("unmatched_model_count") or "—"),
                str(len(row.get("unmatched_estimator_beams") or []) if isinstance(row.get("unmatched_estimator_beams"), list) else row.get("unmatched_gt_count") or "—"),
                str(row.get("truth_source") or "—"),
            ]
        )
    grid(
        doc,
        ["Drawing set", "Model beams", "GT beams", "Matched", "Unmatched model", "Unmatched GT", "Truth source"],
        pop_rows,
        [2.4, 2.1, 1.9, 1.8, 2.4, 2.2, 2.2],
    )
    body(doc, "First Set is excluded from all aggregate KPIs. Unmatched beams are recorded, not silently dropped.", size=9, color=STEEL)

    heading_bar(doc, "4.  Vision execution and hybrid coverage")
    vis_rows = []
    for key in INCLUDED_SET_KEYS:
        row = (vis_exec.get("by_set") or {}).get(key) or {}
        vis_rows.append(
            [
                f"{key} Set",
                str(row.get("eligible") or 0),
                str(row.get("attempted") or 0),
                str(row.get("new_live") or 0),
                str(row.get("reused") or 0),
                str(row.get("retried") or 0),
                str(row.get("api_success") or 0),
                str(row.get("schema_valid") or 0),
                str(row.get("usable") or 0),
                str(row.get("hybrid") or 0),
                str(row.get("fallback") or 0),
            ]
        )
    grid(
        doc,
        ["Set", "Eligible", "Attempted", "New live", "Reused", "Retried", "API ok", "Schema", "Usable", "HYBRID", "FALLBACK"],
        vis_rows,
        [1.7, 1.4, 1.6, 1.4, 1.3, 1.3, 1.3, 1.3, 1.3, 1.4, 1.6],
    )
    body(
        doc,
        f"Fifth Set valid E.2 live Vision artefacts are reused when the current semantic contract, schema, "
        f"source fingerprints and population match ({data.get('fifth_reuse_decision') or 'see artefacts'}). "
        f"Second, Third, Fourth and Sixth use current live Claude (claude-sonnet-4-5) where a valid render exists. "
        f"Individual Vision failures apply the D.1 fallback policy and are labelled FALLBACK, never HYBRID.",
        size=9.5,
    )
    _add_picture(doc, charts.get("hybrid_fallback_coverage.png"))

    heading_bar(doc, "5.  Accuracy by drawing set")
    acc_rows = []
    for key in INCLUDED_SET_KEYS:
        row = per_set.get(key) or {}
        acc_rows.append(
            [
                f"{key} Set",
                _fmt(row.get("beam_identification_percent"), suffix="%"),
                _fmt(row.get("bar_identification_percent"), suffix="%"),
                _fmt(row.get("correct_of_detected_percent"), suffix="%"),
                _fmt(row.get("diameter_identification_percent"), suffix="%"),
                _fmt(row.get("weight_accuracy_percent"), suffix="%"),
                _fmt(row.get("overall_accuracy_percent"), suffix="%"),
            ]
        )
    grid(
        doc,
        ["Drawing set", "Beam ID", "Bar ID", "Correct of detected", "Diameter ID", "Steel", "Overall"],
        acc_rows,
        [2.6, 2.0, 2.0, 2.8, 2.4, 2.0, 2.2],
    )
    _add_picture(doc, charts.get("per_set_overall_accuracy.png"))
    _add_picture(doc, charts.get("per_set_kpi_comparison.png"))

    heading_bar(doc, "6.  Pooled Second-to-Sixth performance")
    grid(
        doc,
        ["KPI", "Pooled result", "Raw counts"],
        [
            ["Beam identification", _fmt(pooled.get("beam_identification_percent"), suffix="%"), f"{pooled.get('beam_n')} / {pooled.get('beam_d')}"],
            ["Bar identification", _fmt(pooled.get("bar_identification_percent"), suffix="%"), f"{pooled.get('bar_n')} / {pooled.get('bar_d')}"],
            ["Correct of detected bars", _fmt(pooled.get("correct_of_detected_percent"), suffix="%"), f"{pooled.get('correct_n')} / {pooled.get('correct_d')}"],
            ["Diameter identification", _fmt(pooled.get("diameter_identification_percent"), suffix="%"), f"{pooled.get('diameter_n')} / {pooled.get('diameter_d')}"],
            ["Steel / weight accuracy", _fmt(pooled.get("weight_accuracy_percent"), suffix="%"), f"{_fmt(pooled.get('hybrid_total_kg'), 3)} / {_fmt(pooled.get('benchmark_total_kg'), 3)} kg"],
            ["Overall accuracy", _fmt(pooled.get("overall_accuracy_percent"), suffix="%"), "Mean of four pooled KPIs; diameter excluded"],
        ],
        [5.0, 3.6, 8.4],
    )
    body(
        doc,
        "Headline pooled results sum raw numerators and denominators. Steel accuracy is computed after pooling kg. "
        "Set percentages are not averaged for the headline overall score.",
        size=9,
        color=STEEL,
    )
    _add_picture(doc, charts.get("model_vs_benchmark_steel_kg.png"))

    heading_bar(doc, "7.  Diameter identification (detected bar lines)")
    body(
        doc,
        "QA.2A’s published field diameter_accuracy_pct is an alias of bar matching accuracy "
        "(MATCH / detected) and is not used here. Diameter identification below is counted from "
        "Bar Matching rows: a detected bar is diameter-correct unless status is WRONG_DIAMETER. "
        "GT diameter is the estimator line diameter. Percentages are pooled from raw counts across "
        "Second–Sixth; set percentages are not averaged. Diameter remains excluded from overall.",
        size=9,
        space_after=5,
    )
    ident_rows = _diameter_ident_rows(data)
    if ident_rows:
        grid(
            doc,
            ["Diameter", "GT bar lines", "Detected", "MATCH", "WRONG_DIA", "Diameter ID", "Note"],
            ident_rows,
            [2.4, 2.2, 1.8, 1.6, 2.0, 2.2, 4.8],
            total_last=True,
        )
    else:
        body(doc, "No diameter-wise identification rows were produced for this report payload.", size=9)

    heading_bar(doc, "8.  Diameter-wise steel quantity  —  not the same as diameter identification")
    body(
        doc,
        "Quantity ratio = automated kg / estimated kg × 100. It is not accuracy. "
        "A ratio above 100% is an overestimate. This table pools kilogram totals by diameter "
        "before computing difference and ratio. It is not diameter identification.",
        size=9,
        space_after=5,
    )
    steel_rows = _diameter_steel_rows(data)
    if steel_rows:
        grid(
            doc,
            ["Diameter", "Estimated kg", "Automated kg", "Difference kg", "Abs % diff", "Quantity ratio"],
            steel_rows,
            [2.4, 3.0, 3.0, 2.8, 2.4, 3.4],
            total_last=True,
        )
    else:
        body(doc, "No diameter-wise steel rows were produced for this report payload.", size=9)

    heading_bar(doc, "9.  Semantic error profile")
    tax = data.get("semantic_taxonomy_pooled") or {}
    tax_rows = [[k, str(v)] for k, v in sorted(tax.items(), key=lambda kv: -int(kv[1] or 0))]
    if tax_rows:
        grid(doc, ["Matcher taxonomy", "Pooled count"], tax_rows, [10.0, 7.0])
    body(
        doc,
        "MISSING bars are ground-truth lines the model did not identify. MATCH / WRONG_QUANTITY / WRONG_DIAMETER / "
        "WRONG_ROLE / PARTIAL_MATCH describe detected bars. EXTRA is a model line without a GT pair. "
        "Deterministic disagreement is not treated as proof that Vision is wrong.",
        size=9.5,
    )
    _add_picture(doc, charts.get("semantic_error_distribution.png"))

    heading_bar(doc, "10.  Engineering calculation profile")
    eng = data.get("engineering_errors") or {}
    counts = eng.get("counts") or eng
    if isinstance(counts, dict):
        grid(
            doc,
            ["Engineering condition", "Count"],
            [[k, str(v)] for k, v in counts.items()],
            [12.0, 5.0],
        )
    body(
        doc,
        "Spacers remain deterministic-only. Stirrup identification may come from Vision; stirrup geometry, legs, "
        "zones, spacing, quantity, hooks, cut length and weight remain deterministic. Vision never invents stirrup "
        "engineering quantities. Semantic/engineering stirrup conflicts are recorded, not silently overwritten.",
        size=9.5,
    )

    heading_bar(doc, "11.  HYBRID / FALLBACK coverage")
    body(
        doc,
        "HYBRID_ONLY scores use the valid matched benchmark subset for Vision-usable beams. "
        "FALLBACK_ONLY uses beams that applied the D.1 deterministic fallback. "
        "Full-population scores are not replaced by HYBRID_ONLY steel accuracy. "
        "A zero cohort is reported as NOT APPLICABLE.",
        size=9.5,
    )
    coh = data.get("cohorts") or {}
    coh_rows = []
    for label in ("HYBRID_ONLY", "FALLBACK_ONLY", "FULL_POPULATION"):
        block = coh.get(label) or {}
        if not block.get("applicable"):
            coh_rows.append([label, "NOT APPLICABLE", "—", "—", "—", "—", "—"])
        else:
            k = block.get("kpis") or block
            coh_rows.append(
                [
                    label,
                    _fmt(k.get("beam_identification_percent"), suffix="%"),
                    _fmt(k.get("bar_identification_percent"), suffix="%"),
                    _fmt(k.get("correct_of_detected_percent"), suffix="%"),
                    _fmt(k.get("diameter_identification_percent"), suffix="%"),
                    _fmt(k.get("weight_accuracy_percent"), suffix="%"),
                    _fmt(k.get("overall_accuracy_percent"), suffix="%"),
                ]
            )
    grid(
        doc,
        ["Cohort", "Beam ID", "Bar ID", "Correct", "Diameter", "Steel", "Overall"],
        coh_rows,
        [3.2, 2.2, 2.2, 2.2, 2.2, 2.2, 2.2],
    )

    heading_bar(doc, "12.  Current workflow value")
    body(
        doc,
        "Positioning: model-assisted estimation. Estimator verification is required. "
        "No current time-and-motion study exists in this phase, so accuracy is not equivalent to time saved. "
        "Vision API success is not engineering accuracy. Fallback coverage is part of the measured system, "
        "not hidden Vision performance.",
        size=10,
    )

    heading_bar(doc, "13.  Methodology, sources and limitations")
    body(
        doc,
        f"{PHASE_ID} uses QA.2A BeamMatcher / BarMatcher / MetricsEngine metric8 and the QA.3.0 four-KPI overall mean. "
        f"Ground truth is evaluation-only (estimator Excel; Second–Third also have a validated QA.2B.1 context path where discovered). "
        f"Truth is never used in runtime semantic resolution. "
        f"Limitations: {', '.join(data.get('limitations') or ['see artefacts'])}. "
        f"PRODUCTION_WRITE = false. ENGINEERING_CHANGES = NONE. Shadow-only; not production routing.",
        size=8.5,
        color=STEEL,
        space_after=2,
    )

    dest = Path(out_root) / DOCX_NAME
    dest.parent.mkdir(parents=True, exist_ok=True)
    doc.save(dest)
    return dest


__all__ = ["write_docx"]
