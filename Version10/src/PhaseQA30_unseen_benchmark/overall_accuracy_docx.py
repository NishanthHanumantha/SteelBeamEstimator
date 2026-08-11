"""
Build Overall Accuracy Word report for all six drawing sets.
Combines QA.2B.1 (First–Third) with QA.3.0 (Fourth–Sixth).
MODEL_VERSION: 10.0.0
"""
from __future__ import annotations

import json
import shutil
from copy import deepcopy
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

from docx import Document
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor
from docx.text.paragraph import Paragraph

MODEL_VERSION = "10.0.0"
PHASE_ID = "QA.3.0"

# Version 9.6.1 three-set baseline (known-set regeneration)
BASELINE_V961 = {
    "beam_detection_pct": 93.92,
    "bar_detection_pct": 68.52,
    "bar_accuracy_pct": 27.41,
    "steel_accuracy_pct": 91.07,
    "overall_accuracy_pct": 70.23,
}

RED = RGBColor(0xEE, 0x00, 0x00)
GREEN = RGBColor(0x00, 0x80, 0x00)


def _pct(v: float, digits: int = 2) -> str:
    return f"{float(v):.{digits}f}%"


def _delta(curr: float, base: float) -> str:
    d = float(curr) - float(base)
    sign = "+" if d >= 0 else ""
    return f"{sign}{d:.2f}%"


def _clear_paragraph(para: Paragraph) -> None:
    for run in list(para.runs):
        run._element.getparent().remove(run._element)


def _heading(para: Paragraph, text: str, *, size_pt: Optional[float] = None) -> None:
    _clear_paragraph(para)
    run = para.add_run(text)
    run.bold = True
    if size_pt is not None:
        run.font.size = Pt(size_pt)


def _label_value(para: Paragraph, label: str, value: str) -> None:
    _clear_paragraph(para)
    para.add_run(label)
    r = para.add_run(value)
    r.bold = True


def _metric_line(para: Paragraph, label: str, value: str) -> None:
    _clear_paragraph(para)
    r0 = para.add_run(f"{label} ")
    r0.bold = True
    r1 = para.add_run(value)
    r1.bold = True


def _plain(para: Paragraph, text: str) -> None:
    _clear_paragraph(para)
    para.add_run(text)


def _insert_paragraph_after(paragraph: Paragraph, text: str = "", *, bold: bool = False) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if text:
        run = new_para.add_run(text)
        run.bold = bold
    return new_para


def _insert_paragraph_before(paragraph: Paragraph, text: str = "", *, bold: bool = False) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addprevious(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if text:
        run = new_para.add_run(text)
        run.bold = bold
    return new_para


def _insert_table_after(paragraph: Paragraph, rows: int, cols: int):
    doc = paragraph.part.document
    table = doc.add_table(rows=rows, cols=cols)
    tbl = table._tbl
    paragraph._p.addnext(tbl)
    return table


def _set_cell(cell, text: str, *, bold: bool = False, color: Optional[RGBColor] = None) -> None:
    para = cell.paragraphs[0]
    _clear_paragraph(para)
    run = para.add_run(text)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def _find_para(doc: Document, exact: str = "", *, startswith: str = "") -> Paragraph:
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if exact and t == exact:
            return p
        if startswith and t.startswith(startswith):
            return p
    raise KeyError(exact or startswith)


def _find_para_contains(doc: Document, needle: str) -> Paragraph:
    for p in doc.paragraphs:
        if needle in (p.text or ""):
            return p
    raise KeyError(needle)


def _ensure_table_cols(table, cols: int) -> None:
    # python-docx may not refresh table.columns immediately after appending cells,
    # so never loop on len(table.columns). Add a fixed number of columns once.
    current = len(table.rows[0].cells) if table.rows else 0
    missing = max(0, cols - current)
    for _ in range(missing):
        for row in table.rows:
            row._tr.append(deepcopy(row.cells[-1]._tc))


def _ensure_table_rows(table, rows: int) -> None:
    missing = max(0, rows - len(table.rows))
    for _ in range(missing):
        table.add_row()


def _overall_of(beam: float, bar_det: float, bar_acc: float, steel: float) -> float:
    return round((beam + bar_det + bar_acc + steel) / 4.0, 2)


def load_six_set_bundle(
    *,
    qa2b1_json: Path,
    qa30_json: Path,
    qa30_root: Path,
) -> Dict[str, Any]:
    known = json.loads(Path(qa2b1_json).read_text(encoding="utf-8"))
    unseen = json.loads(Path(qa30_json).read_text(encoding="utf-8"))

    sets: List[Dict[str, Any]] = []

    # First–Third from QA.2B.1
    dash = {
        (r.get("drawing_set") or ""): r
        for r in ((known.get("dashboard") or {}).get("drawing_sets") or [])
    }
    for s in known.get("drawing_sets") or []:
        name = s["drawing_set"]
        bm = s.get("beam_matching") or {}
        bar = s.get("bar_matching") or {}
        st = s.get("steel") or {}
        est = s.get("estimator_summary") or {}
        drow = dash.get(name) or {}
        beam_det = float(bm.get("detection_pct") or 0)
        bar_det = float(bar.get("detection_pct") or 0)
        bar_acc = float(bar.get("accuracy_pct") or 0)
        steel = float(st.get("accuracy_pct") or 0)
        est_bars = int(est.get("bars") or 0)
        miss = int(bar.get("missing_bars") or 0)
        det_bars = max(0, est_bars - miss) if est_bars else int(round(est_bars * bar_det / 100.0))
        # Prefer count consistency with published totals when possible
        if est_bars and bar_det:
            det_bars = int(round(est_bars * bar_det / 100.0))
        corr_bars = int(round(det_bars * bar_acc / 100.0)) if det_bars else 0
        sets.append(
            {
                "drawing_set": name,
                "group": "known",
                "source": "QA.2B.1",
                "beam_detection_pct": beam_det,
                "bar_detection_pct": bar_det,
                "bar_accuracy_pct": bar_acc,
                "steel_accuracy_pct": steel,
                "overall_accuracy_pct": float(
                    drow.get("overall_accuracy_pct")
                    or _overall_of(beam_det, bar_det, bar_acc, steel)
                ),
                "estimator_beams": int(bm.get("estimator_beams") or est.get("beams") or 0),
                "detected_beams": int(bm.get("detected_beams") or 0),
                "estimator_bars": est_bars,
                "detected_bars": det_bars,
                "correct_bars": corr_bars,
                "missing_bars": miss,
                "estimator_kg": float(est.get("kg") or st.get("estimator_total_kg") or 0),
                "model_kg": float((s.get("model_summary") or {}).get("kg") or st.get("model_total_kg") or 0),
                "error_count": drow.get("error_count"),
            }
        )

    # Fourth–Sixth from QA.3.0 per-set artefacts + summary
    for row in unseen.get("drawing_set_results") or []:
        name = row["drawing_set"]
        key = name.replace(" ", "_")
        br_path = Path(qa30_root) / key / "benchmark_result.json"
        br = json.loads(br_path.read_text(encoding="utf-8")) if br_path.exists() else {}
        bm = br.get("beam_matching") or {}
        bar = br.get("bar_matching") or {}
        st = br.get("steel") or {}
        est = br.get("estimator_summary") or {}
        beam_det = float(row.get("beam_detection_pct") or bm.get("detection_pct") or 0)
        bar_det = float(row.get("bar_detection_pct") or bar.get("detection_pct") or 0)
        bar_acc = float(row.get("bar_accuracy_pct") or bar.get("accuracy_pct") or 0)
        steel = float(row.get("steel_accuracy_pct") or st.get("accuracy_pct") or 0)
        est_bars = int(est.get("bars") or 0)
        miss = int(bar.get("missing_bars") or 0)
        det_bars = int(round(est_bars * bar_det / 100.0)) if est_bars else 0
        corr_bars = int(round(det_bars * bar_acc / 100.0)) if det_bars else 0
        sets.append(
            {
                "drawing_set": name,
                "group": "unseen",
                "source": "QA.3.0",
                "beam_detection_pct": beam_det,
                "bar_detection_pct": bar_det,
                "bar_accuracy_pct": bar_acc,
                "steel_accuracy_pct": steel,
                "overall_accuracy_pct": float(
                    row.get("overall_accuracy_pct")
                    or _overall_of(beam_det, bar_det, bar_acc, steel)
                ),
                "estimator_beams": int(bm.get("estimator_beams") or est.get("beams") or 0),
                "detected_beams": int(bm.get("detected_beams") or 0),
                "estimator_bars": est_bars,
                "detected_bars": det_bars,
                "correct_bars": corr_bars,
                "missing_bars": miss,
                "estimator_kg": float(est.get("kg") or st.get("estimator_total_kg") or 0),
                "model_kg": float((br.get("model_summary") or {}).get("kg") or st.get("model_total_kg") or 0),
                "error_count": row.get("error_count"),
            }
        )

    # Prefer published aggregate counts where available, else sum
    known_bench = known.get("benchmark") or {}
    unseen_bench = unseen.get("benchmark") or {}
    total_beams = int(known_bench.get("total_beams") or 0) + int(unseen_bench.get("total_beams") or 0)
    detected_beams = int(known_bench.get("detected_beams") or 0) + int(
        unseen_bench.get("detected_beams") or 0
    )
    total_bars = int(known_bench.get("total_bars") or 0) + int(unseen_bench.get("total_bars") or 0)
    detected_bars = int(known_bench.get("detected_bars") or 0) + int(
        unseen_bench.get("detected_bars") or 0
    )
    correct_bars = int(known_bench.get("correct_bars") or 0) + int(
        unseen_bench.get("correct_bars") or 0
    )
    missing_bars = int(known_bench.get("missing_bars") or 0) + int(
        unseen_bench.get("missing_bars") or 0
    )
    if not total_beams:
        total_beams = sum(s["estimator_beams"] for s in sets)
        detected_beams = sum(s["detected_beams"] for s in sets)
    if not total_bars:
        total_bars = sum(s["estimator_bars"] for s in sets)
        detected_bars = sum(s["detected_bars"] for s in sets)
        correct_bars = sum(s["correct_bars"] for s in sets)
        missing_bars = sum(s["missing_bars"] for s in sets)

    beam_det = round(100.0 * detected_beams / total_beams, 2) if total_beams else 0.0
    bar_det = round(100.0 * detected_bars / total_bars, 2) if total_bars else 0.0
    bar_acc = round(100.0 * correct_bars / detected_bars, 2) if detected_bars else 0.0
    steel_acc = round(sum(s["steel_accuracy_pct"] for s in sets) / len(sets), 2) if sets else 0.0
    overall = _overall_of(beam_det, bar_det, bar_acc, steel_acc)

    # Merge error frequencies
    freq: Dict[str, int] = {}
    for src in (known.get("errors") or {}, unseen.get("errors") or {}):
        for k, v in (src.get("frequency") or {}).items():
            freq[k] = freq.get(k, 0) + int(v)

    known_group = [s for s in sets if s["group"] == "known"]
    unseen_group = [s for s in sets if s["group"] == "unseen"]

    def _group_metrics(rows: List[Dict[str, Any]]) -> Dict[str, float]:
        if not rows:
            return {}
        tb = sum(r["estimator_beams"] for r in rows)
        db = sum(r["detected_beams"] for r in rows)
        tar = sum(r["estimator_bars"] for r in rows)
        dar = sum(r["detected_bars"] for r in rows)
        car = sum(r["correct_bars"] for r in rows)
        # Prefer published group metrics when exact
        if rows[0]["group"] == "known":
            return {
                "beam_detection_pct": float(known_bench.get("beam_detection_pct") or 0),
                "bar_detection_pct": float(known_bench.get("bar_detection_pct") or 0),
                "bar_accuracy_pct": float(known_bench.get("bar_accuracy_pct") or 0),
                "steel_accuracy_pct": float(known_bench.get("steel_accuracy_pct") or 0),
                "overall_accuracy_pct": float(known_bench.get("overall_accuracy_pct") or 0),
            }
        om = unseen.get("overall_metrics") or {}
        return {
            "beam_detection_pct": float(om.get("beam_detection_pct") or unseen_bench.get("beam_detection_pct") or 0),
            "bar_detection_pct": float(om.get("bar_detection_pct") or unseen_bench.get("bar_detection_pct") or 0),
            "bar_accuracy_pct": float(
                om.get("bar_matching_pct") or om.get("bar_accuracy_pct") or unseen_bench.get("bar_accuracy_pct") or 0
            ),
            "steel_accuracy_pct": float(om.get("steel_accuracy_pct") or unseen_bench.get("steel_accuracy_pct") or 0),
            "overall_accuracy_pct": float(om.get("overall_accuracy_pct") or unseen_bench.get("overall_accuracy_pct") or 0),
        }

    return {
        "sets": sets,
        "combined": {
            "total_drawing_sets": len(sets),
            "total_beams": total_beams,
            "detected_beams": detected_beams,
            "correct_beams": detected_beams,
            "total_bars": total_bars,
            "detected_bars": detected_bars,
            "correct_bars": correct_bars,
            "missing_bars": missing_bars,
            "beam_detection_pct": beam_det,
            "bar_detection_pct": bar_det,
            "bar_accuracy_pct": bar_acc,
            "steel_accuracy_pct": steel_acc,
            "overall_accuracy_pct": overall,
        },
        "known_metrics": _group_metrics(known_group),
        "unseen_metrics": _group_metrics(unseen_group),
        "errors": freq,
        "sources": {
            "qa2b1": str(qa2b1_json),
            "qa30": str(qa30_json),
        },
    }


def build_report(
    *,
    template_path: Path,
    qa2b1_json: Path,
    qa30_json: Path,
    qa30_root: Path,
    output_docx: Path,
) -> Dict[str, Any]:
    template_path = Path(template_path)
    output_docx = Path(output_docx)
    output_docx.parent.mkdir(parents=True, exist_ok=True)

    bundle = load_six_set_bundle(
        qa2b1_json=qa2b1_json, qa30_json=qa30_json, qa30_root=qa30_root
    )
    sets = {s["drawing_set"]: s for s in bundle["sets"]}
    bench = bundle["combined"]
    errors = bundle["errors"]
    curr = {
        "beam_detection_pct": bench["beam_detection_pct"],
        "bar_detection_pct": bench["bar_detection_pct"],
        "bar_accuracy_pct": bench["bar_accuracy_pct"],
        "steel_accuracy_pct": bench["steel_accuracy_pct"],
        "overall_accuracy_pct": bench["overall_accuracy_pct"],
    }

    shutil.copy2(template_path, output_docx)
    doc = Document(str(output_docx))

    # Title
    _heading(
        doc.paragraphs[0],
        "Steel Beam Estimation – Ground Truth Benchmark Summary",
        size_pt=14,
    )
    _heading(
        doc.paragraphs[1],
        "Version 10.0.0 | All Six Drawing Sets — Known (QA.2B.1) + Unseen (QA.3.0)",
    )

    _heading(doc.paragraphs[2], "Benchmark Results")
    _plain(
        doc.paragraphs[3],
        "This report consolidates ground-truth accuracy across all six drawing sets. "
        "First–Third Set results are taken from Version 9 QA.2B.1 regenerated production outputs "
        "(known / development drawing sets). Fourth–Sixth Set results are taken from Version 10 "
        "QA.3.0 unseen generalization benchmark production outputs. For every set, model Estimation_Output.xlsx "
        "was compared against the estimator bar bending schedule (ground truth only). "
        "Estimator Excel was never used during DXF production.",
    )

    # Comparison vs V9.6.1 known-set baseline
    dw = _find_para(doc, "Drawing-wise Performance")
    h_imp = _insert_paragraph_before(dw, "Accuracy Summary vs Version 9.6.1 (First–Third Baseline)", bold=True)
    intro = _insert_paragraph_after(
        h_imp,
        "The table below compares the Version 9.6.1 three-set (First–Third) baseline against the "
        "combined six-set accuracy (First–Sixth).",
    )
    imp_rows = [
        ("Metric", "V9.6.1 First–Third", "V10 All Six Sets", "Delta"),
        (
            "Beam Detection",
            _pct(BASELINE_V961["beam_detection_pct"]),
            _pct(curr["beam_detection_pct"]),
            _delta(curr["beam_detection_pct"], BASELINE_V961["beam_detection_pct"]),
        ),
        (
            "Bar Detection",
            _pct(BASELINE_V961["bar_detection_pct"]),
            _pct(curr["bar_detection_pct"]),
            _delta(curr["bar_detection_pct"], BASELINE_V961["bar_detection_pct"]),
        ),
        (
            "Bar Matching Accuracy",
            _pct(BASELINE_V961["bar_accuracy_pct"]),
            _pct(curr["bar_accuracy_pct"]),
            _delta(curr["bar_accuracy_pct"], BASELINE_V961["bar_accuracy_pct"]),
        ),
        (
            "Steel Accuracy",
            _pct(BASELINE_V961["steel_accuracy_pct"]),
            _pct(curr["steel_accuracy_pct"]),
            _delta(curr["steel_accuracy_pct"], BASELINE_V961["steel_accuracy_pct"]),
        ),
        (
            "Overall Accuracy",
            _pct(BASELINE_V961["overall_accuracy_pct"]),
            _pct(curr["overall_accuracy_pct"]),
            _delta(curr["overall_accuracy_pct"], BASELINE_V961["overall_accuracy_pct"]),
        ),
    ]
    imp_table = _insert_table_after(intro, len(imp_rows), 4)
    for t in doc.tables:
        if t.rows and "Drawing Set" in (t.rows[0].cells[0].text or ""):
            try:
                imp_table.style = t.style
            except Exception:
                pass
            break
    for ri, row_vals in enumerate(imp_rows):
        for ci, val in enumerate(row_vals):
            color = None
            if ri > 0 and ci == 3:
                if val.startswith("+") and not val.startswith("+0.00"):
                    color = GREEN
                elif val.startswith("-"):
                    color = RED
            _set_cell(imp_table.rows[ri].cells[ci], val, bold=(ri == 0 or ci == 0), color=color)
    _insert_paragraph_after(Paragraph(imp_table._tbl, intro._parent), "")

    # Known vs Unseen subgroup table
    sub_h = _insert_paragraph_after(
        Paragraph(imp_table._tbl, intro._parent),
        "Known vs Unseen Subgroup Accuracy",
        bold=True,
    )
    sub_intro = _insert_paragraph_after(
        sub_h,
        "First–Third are known development sets (QA.2B.1). Fourth–Sixth are unseen generalization sets (QA.3.0).",
    )
    km = bundle["known_metrics"]
    um = bundle["unseen_metrics"]
    sub_rows = [
        ("Metric", "First–Third (Known)", "Fourth–Sixth (Unseen)", "All Six Sets"),
        ("Beam Detection", _pct(km["beam_detection_pct"]), _pct(um["beam_detection_pct"]), _pct(curr["beam_detection_pct"])),
        ("Bar Detection", _pct(km["bar_detection_pct"]), _pct(um["bar_detection_pct"]), _pct(curr["bar_detection_pct"])),
        ("Bar Matching", _pct(km["bar_accuracy_pct"]), _pct(um["bar_accuracy_pct"]), _pct(curr["bar_accuracy_pct"])),
        ("Steel Accuracy", _pct(km["steel_accuracy_pct"]), _pct(um["steel_accuracy_pct"]), _pct(curr["steel_accuracy_pct"])),
        ("Overall Accuracy", _pct(km["overall_accuracy_pct"]), _pct(um["overall_accuracy_pct"]), _pct(curr["overall_accuracy_pct"])),
    ]
    sub_table = _insert_table_after(sub_intro, len(sub_rows), 4)
    try:
        sub_table.style = imp_table.style
    except Exception:
        pass
    for ri, row_vals in enumerate(sub_rows):
        for ci, val in enumerate(row_vals):
            _set_cell(sub_table.rows[ri].cells[ci], val, bold=(ri == 0 or ci == 0))
    _insert_paragraph_after(Paragraph(sub_table._tbl, sub_intro._parent), "")

    # Drawing-wise Performance — replace template table with a fresh 6-set table
    dw_heading = _find_para(doc, "Drawing-wise Performance")
    _heading(dw_heading, "Drawing-wise Performance")

    # Clear legacy template drawing-wise table cells (keep structure, avoid col expansion bugs)
    legacy = None
    for t in doc.tables:
        if t.rows and "Drawing Set" in (t.rows[0].cells[0].text or ""):
            legacy = t
            break
    if legacy is not None:
        for row in legacy.rows:
            for cell in row.cells:
                _set_cell(cell, "")

    order = [
        ("First Set", "First Set Drawings"),
        ("Second Set", "Second Set Drawings"),
        ("Third Set", "Third Set Drawings"),
        ("Fourth Set", "Fourth Set Drawings"),
        ("Fifth Set", "Fifth Set Drawings"),
        ("Sixth Set", "Sixth Set Drawings"),
    ]

    def beam_cell(ds: Dict[str, Any]) -> str:
        return (
            f"{ds['detected_beams']} / {ds['estimator_beams']} "
            f"({ds['beam_detection_pct']:.1f}%)"
        )

    draw_rows = [
        [
            "Drawing Set",
            "Beam Detection",
            "Bar Detection",
            "Bar Accuracy",
            "Steel Accuracy",
            "Overall",
        ]
    ]
    for label, key in order:
        ds = sets[key]
        draw_rows.append(
            [
                label,
                beam_cell(ds),
                _pct(ds["bar_detection_pct"]),
                _pct(ds["bar_accuracy_pct"]),
                _pct(ds["steel_accuracy_pct"]),
                _pct(ds["overall_accuracy_pct"]),
            ]
        )
    draw_rows.append(
        [
            "All Six Sets",
            f"{bench['detected_beams']} / {bench['total_beams']} ({bench['beam_detection_pct']:.2f}%)",
            _pct(bench["bar_detection_pct"]),
            _pct(bench["bar_accuracy_pct"]),
            _pct(bench["steel_accuracy_pct"]),
            _pct(bench["overall_accuracy_pct"]),
        ]
    )
    t0 = _insert_table_after(dw_heading, len(draw_rows), 6)
    try:
        if legacy is not None:
            t0.style = legacy.style
    except Exception:
        pass
    for ri, vals in enumerate(draw_rows):
        for ci, val in enumerate(vals):
            color = None
            if ri > 0 and ci in (3, 5):
                try:
                    if float(str(val).rstrip("%")) < 50:
                        color = RED
                except ValueError:
                    pass
            _set_cell(
                t0.rows[ri].cells[ci],
                val,
                bold=(ri == 0 or ri == len(draw_rows) - 1),
                color=color,
            )

    try:
        narr = _find_para_contains(doc, "consistent beam detection")
        _plain(
            narr,
            "Across all six sets, beam detection remains relatively strong, but drops on unseen "
            "Fourth/Fifth projects. Steel accuracy is high on First and Sixth sets and weaker on "
            "Fourth/Fifth. Reinforcement bar detection and matching remain the primary limiters of "
            "overall estimation accuracy, especially under unseen generalization conditions.",
        )
    except KeyError:
        pass

    # Overall Performance Summary
    _heading(_find_para(doc, "Overall Performance Summary"), "Overall Performance Summary")
    _heading(_find_para(doc, "Beam Detection"), "Beam Detection")
    _label_value(
        _find_para(doc, startswith="Ground Truth Beams:"),
        "Ground Truth Beams: ",
        str(bench["total_beams"]),
    )
    _label_value(
        _find_para(doc, startswith="Correctly Detected"),
        "Correctly Detected Beams: ",
        str(bench["detected_beams"]),
    )
    _metric_line(
        _find_para(doc, startswith="Beam Detection Accuracy"),
        "Beam Detection Accuracy:",
        _pct(bench["beam_detection_pct"]),
    )
    _heading(
        _find_para(doc, "Reinforcement Bar Interpretation"),
        "Reinforcement Bar Interpretation",
    )
    _label_value(
        _find_para(doc, startswith="Ground Truth Bars:"),
        "Ground Truth Bars: ",
        str(bench["total_bars"]),
    )
    _label_value(
        _find_para(doc, startswith="Detected Bars:"),
        "Detected Bars: ",
        str(bench["detected_bars"]),
    )
    _label_value(
        _find_para(doc, startswith="Correctly Matched Bars:"),
        "Correctly Matched Bars: ",
        str(bench["correct_bars"]),
    )
    _label_value(
        _find_para(doc, startswith="Missing Bars:"),
        "Missing Bars: ",
        str(bench["missing_bars"]),
    )
    _metric_line(
        _find_para(doc, startswith="Bar Detection Accuracy"),
        "Bar Detection Accuracy:",
        _pct(bench["bar_detection_pct"]),
    )
    _metric_line(
        _find_para(doc, startswith="Bar Matching Accuracy"),
        "Bar Matching Accuracy:",
        _pct(bench["bar_accuracy_pct"]),
    )
    _heading(_find_para(doc, "Steel Quantity Estimation"), "Steel Quantity Estimation")
    _plain(
        _find_para(doc, startswith="Average Steel Quantity Accuracy"),
        f"Average Steel Quantity Accuracy: {_pct(bench['steel_accuracy_pct'])}",
    )
    _heading(_find_para(doc, "Overall Model Accuracy"), "Overall Model Accuracy")
    _plain(
        _find_para(doc, startswith="Overall Ground Truth Benchmark Accuracy"),
        f"Overall Ground Truth Benchmark Accuracy (All Six Sets): {_pct(bench['overall_accuracy_pct'])}",
    )

    # Key Error Categories
    _heading(_find_para(doc, "Key Error Categories"), "Key Error Categories")
    try:
        _plain(
            _find_para_contains(doc, "automatically classified discrepancies"),
            "The benchmark automatically classified discrepancies between regenerated model outputs "
            "and estimator ground truth across all six drawing sets (combined QA.2B.1 + QA.3.0 error tallies).",
        )
    except KeyError:
        pass

    err_table = None
    for t in doc.tables:
        if t.rows and "Error" in (t.rows[0].cells[0].text or ""):
            err_table = t
            break
    assert err_table is not None
    err_rows = [
        ("Error Category", "Count"),
        ("Missing Bars", str(errors.get("Missing Bar", 0))),
        ("Wrong Quantity", str(errors.get("Wrong Quantity", 0))),
        ("Extra Bars", str(errors.get("Extra Bar", 0))),
        ("Wrong Diameter", str(errors.get("Wrong Diameter", 0))),
        ("Wrong Role Classification", str(errors.get("Wrong Role", 0))),
        ("Beam Missing", str(errors.get("Beam Missing", 0))),
        ("Beam Mismatch", str(errors.get("Beam Mismatch", 0))),
        ("Steel Difference", str(errors.get("Steel Difference", 0))),
    ]
    _ensure_table_rows(err_table, len(err_rows))
    for ri, vals in enumerate(err_rows):
        _set_cell(err_table.rows[ri].cells[0], vals[0], bold=(ri == 0))
        _set_cell(err_table.rows[ri].cells[1], vals[1], bold=(ri == 0))
    for ri in range(len(err_rows), len(err_table.rows)):
        _set_cell(err_table.rows[ri].cells[0], "")
        _set_cell(err_table.rows[ri].cells[1], "")

    try:
        p = _find_para_contains(doc, "reinforcement interpretation")
        if "Most of" in p.text or "errors are associated" in p.text:
            _clear_paragraph(p)
            p.add_run("Most of the errors remain associated with ")
            rb = p.add_run("reinforcement interpretation")
            rb.bold = True
            p.add_run(
                ", particularly missing bars, wrong quantity, extra bars, and diameter resolution. "
                f"Total classified errors (all six sets): {sum(int(v) for v in errors.values())}."
            )
    except KeyError:
        pass

    # Engineering context
    kf = _find_para(doc, "Key Findings")
    eng_h = _insert_paragraph_before(
        kf, "Benchmark Scope and Engineering Context", bold=True
    )
    eng_intro = _insert_paragraph_after(
        eng_h,
        "This six-set report spans the Version 9 known-set regeneration programme and the "
        "Version 10 unseen generalization benchmark:",
    )
    prev = eng_intro
    for item in [
        "First–Third Set: QA.2B.1 production regeneration (known development drawings)",
        "Fourth–Sixth Set: QA.3.0 unseen generalization benchmark",
        "DXF-only production; estimator Excel used only for post-run GT comparison",
        "Beam Ownership / Annotation Graph / Adaptive Extent / Shared Scope (Track 1)",
        "Stirrup recovery and geometric evidence pipeline",
        "P2.1–P2.4 leader-chain / bar-failure diagnostic programme (no formula change in this report)",
    ]:
        prev = _insert_paragraph_after(prev, f"• {item}")
    _insert_paragraph_after(prev, "")

    # Key Findings
    kf = _find_para(doc, "Key Findings")
    _heading(kf, "Key Findings")
    findings = [
        f"Combined six-set overall accuracy is {_pct(bench['overall_accuracy_pct'])} "
        f"(beam {_pct(bench['beam_detection_pct'])}, bar detection {_pct(bench['bar_detection_pct'])}, "
        f"bar matching {_pct(bench['bar_accuracy_pct'])}, steel {_pct(bench['steel_accuracy_pct'])}).",
        f"Known sets (First–Third) remain stronger overall at {_pct(km['overall_accuracy_pct'])}, "
        f"while unseen sets (Fourth–Sixth) achieve {_pct(um['overall_accuracy_pct'])}.",
        f"Unseen generalization reduces beam detection from {_pct(km['beam_detection_pct'])} "
        f"to {_pct(um['beam_detection_pct'])} and steel accuracy from {_pct(km['steel_accuracy_pct'])} "
        f"to {_pct(um['steel_accuracy_pct'])}.",
        f"Sixth Set is the strongest unseen project (overall {_pct(sets['Sixth Set Drawings']['overall_accuracy_pct'])}, "
        f"steel {_pct(sets['Sixth Set Drawings']['steel_accuracy_pct'])}).",
        f"Fourth and Fifth Sets are the weakest unseen projects "
        f"(overall {_pct(sets['Fourth Set Drawings']['overall_accuracy_pct'])} and "
        f"{_pct(sets['Fifth Set Drawings']['overall_accuracy_pct'])}).",
        "Reinforcement interpretation — especially missing bars, quantity, and diameter — "
        "remains the dominant accuracy limiter across all six sets.",
        f"Versus the Version 9.6.1 First–Third baseline overall of {_pct(BASELINE_V961['overall_accuracy_pct'])}, "
        f"the all-six-sets overall is {_pct(bench['overall_accuracy_pct'])} "
        f"({_delta(bench['overall_accuracy_pct'], BASELINE_V961['overall_accuracy_pct'])}).",
    ]
    paras = list(doc.paragraphs)
    kf_i = next(i for i, p in enumerate(paras) if (p.text or "").strip() == "Key Findings")
    tail = paras[kf_i + 1 :]
    for i, text in enumerate(findings):
        if i < len(tail):
            _plain(tail[i], text)
        else:
            prev = tail[-1] if tail else kf
            np = _insert_paragraph_after(prev, text)
            tail.append(np)
    for p in tail[len(findings) :]:
        _clear_paragraph(p)

    doc.save(str(output_docx))

    meta = {
        "phase_id": PHASE_ID,
        "model_version": MODEL_VERSION,
        "generated_at": datetime.now(timezone.utc).isoformat(),
        "template": str(template_path),
        "output": str(output_docx),
        "baseline_v961_first_third": BASELINE_V961,
        "combined_all_six": curr,
        "known_first_third": km,
        "unseen_fourth_sixth": um,
        "delta_vs_v961": {k: round(curr[k] - BASELINE_V961[k], 2) for k in curr},
        "sets": bundle["sets"],
        "combined_counts": {
            k: bench[k]
            for k in (
                "total_drawing_sets",
                "total_beams",
                "detected_beams",
                "total_bars",
                "detected_bars",
                "correct_bars",
                "missing_bars",
            )
        },
        "errors": errors,
        "sources": bundle["sources"],
    }
    output_docx.with_suffix(".meta.json").write_text(
        json.dumps(meta, indent=2), encoding="utf-8"
    )

    # Companion markdown summary
    md_path = output_docx.with_suffix(".md")
    md_lines = [
        "# Overall Accuracy Report — All Six Drawing Sets",
        "",
        f"- MODEL_VERSION: `{MODEL_VERSION}`",
        f"- Generated: `{meta['generated_at']}`",
        "",
        "## All Six Sets Combined",
        "",
        f"| Metric | Value |",
        f"|--------|------:|",
        f"| Beam Detection | {_pct(curr['beam_detection_pct'])} |",
        f"| Bar Detection | {_pct(curr['bar_detection_pct'])} |",
        f"| Bar Matching | {_pct(curr['bar_accuracy_pct'])} |",
        f"| Steel Accuracy | {_pct(curr['steel_accuracy_pct'])} |",
        f"| Overall Accuracy | {_pct(curr['overall_accuracy_pct'])} |",
        "",
        "## Drawing-wise",
        "",
        "| Set | Beam | Bar Det | Bar Match | Steel | Overall |",
        "|-----|-----:|--------:|----------:|------:|--------:|",
    ]
    for label, key in order:
        ds = sets[key]
        md_lines.append(
            f"| {label} | {_pct(ds['beam_detection_pct'])} | {_pct(ds['bar_detection_pct'])} | "
            f"{_pct(ds['bar_accuracy_pct'])} | {_pct(ds['steel_accuracy_pct'])} | "
            f"{_pct(ds['overall_accuracy_pct'])} |"
        )
    md_lines += [
        "",
        f"Word report: `{output_docx}`",
        "",
    ]
    md_path.write_text("\n".join(md_lines), encoding="utf-8")
    meta["markdown"] = str(md_path)
    output_docx.with_suffix(".meta.json").write_text(
        json.dumps(meta, indent=2), encoding="utf-8"
    )
    return meta
