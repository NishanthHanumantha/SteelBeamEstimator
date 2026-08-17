"""Stakeholder report: Second–Sixth Sets. Reporting only — no production changes."""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

NAVY = RGBColor(0x1B, 0x36, 0x5D)
STEEL = RGBColor(0x2F, 0x5D, 0x7C)
GREEN = RGBColor(0x1F, 0x6B, 0x3A)
RED = RGBColor(0x9B, 0x2C, 0x2C)
AMBER = RGBColor(0x8A, 0x5A, 0x00)
SLATE = RGBColor(0x3D, 0x4A, 0x57)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
ROW_ALT = "F4F7FA"
HEADER_FILL = "1B365D"
ACCENT_FILL = "2F5D7C"
GREEN_FILL = "1F6B3A"
RED_FILL = "9B2C2C"
AMBER_FILL = "F4E6C4"
GOLD_FILL = "B88A2E"

OUT_DIR = Path(__file__).resolve().parent
DOCX = OUT_DIR / "Steel_Beam_Estimation_Performance_Report_Second_to_Sixth_Sets.docx"
PDF = OUT_DIR / "Steel_Beam_Estimation_Performance_Report_Second_to_Sixth_Sets.pdf"


def _set_run(run, *, size=11, bold=False, color=SLATE, font="Calibri"):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = color


def _shade(cell, hex_color: str) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def _borders(cell, color="C5CDD6") -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        tcBorders.append(el)
    tcPr.append(tcBorders)
    tcMar = OxmlElement("w:tcMar")
    for edge, val in (("left", "60"), ("right", "60"), ("top", "30"), ("bottom", "30")):
        node = OxmlElement(f"w:{edge}")
        node.set(qn("w:w"), val)
        node.set(qn("w:type"), "dxa")
        tcMar.append(node)
    tcPr.append(tcMar)


def cell_text(cell, text, *, bold=False, color=SLATE, size=9, align="center", fill=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
    }[align]
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run(str(text))
    _set_run(run, size=size, bold=bold, color=color)
    if fill:
        _shade(cell, fill)
    _borders(cell)


def set_widths(table, widths_cm):
    table.autofit = False
    table.allow_autofit = False
    total = sum(widths_cm)
    tblPr = table._tbl.tblPr if table._tbl.tblPr is not None else OxmlElement("w:tblPr")
    tblW = tblPr.find(qn("w:tblW"))
    if tblW is None:
        tblW = OxmlElement("w:tblW")
        tblPr.append(tblW)
    tblW.set(qn("w:w"), str(int(total * 567)))
    tblW.set(qn("w:type"), "dxa")
    grid = table._tbl.find(qn("w:tblGrid"))
    if grid is not None:
        for i, w in enumerate(widths_cm):
            list(grid)[i].set(qn("w:w"), str(int(w * 567)))
    for row in table.rows:
        for i, w in enumerate(widths_cm):
            row.cells[i].width = Cm(w)


def heading_bar(doc, title: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell_text(table.cell(0, 0), title, bold=True, color=WHITE, size=11, fill=HEADER_FILL)
    set_widths(table, [17.0])
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(3)


def grid(doc, headers, data, col_widths, *, header_fill=HEADER_FILL, total_last=False):
    table = doc.add_table(rows=1 + len(data), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(headers):
        cell_text(table.cell(0, j), h, bold=True, color=WHITE, size=8, fill=header_fill)
    for i, row in enumerate(data):
        is_total = total_last and i == len(data) - 1
        fill = "E8EEF4" if is_total else (ROW_ALT if i % 2 else "FFFFFF")
        for j, val in enumerate(row):
            cell_text(
                table.cell(i + 1, j),
                val,
                bold=is_total or j == 0,
                color=NAVY if j == 0 else SLATE,
                size=8.5,
                align="left" if j == 0 else "center",
                fill=fill,
            )
    set_widths(table, col_widths)
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(6)


def body(doc, text, *, size=10, color=SLATE, space_after=6, bold=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing = 1.08
    run = p.add_run(text)
    _set_run(run, size=size, color=color, bold=bold)
    return p


def kpis(doc, items):
    table = doc.add_table(rows=2, cols=len(items))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    fills = [HEADER_FILL, ACCENT_FILL, GREEN_FILL, GOLD_FILL]
    for i, (label, value) in enumerate(items):
        cell_text(table.cell(0, i), label, bold=True, color=WHITE, size=7.5, fill=fills[i % len(fills)])
        cell_text(table.cell(1, i), value, bold=True, color=NAVY, size=13, fill="FFFFFF")
    set_widths(table, [17.0 / len(items)] * len(items))
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(8)


def build() -> Path:
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Cm(21.0)
    sec.page_height = Cm(29.7)
    sec.left_margin = Cm(1.7)
    sec.right_margin = Cm(1.7)
    sec.top_margin = Cm(1.3)
    sec.bottom_margin = Cm(1.5)

    hp = sec.header.paragraphs[0]
    r = hp.add_run("SteelBeam Estimator  ·  Version 10  ·  Deterministic production authority")
    _set_run(r, size=8, bold=True, color=STEEL)
    r2 = hp.add_run("     Second–Sixth Sets  ·  First Set excluded")
    _set_run(r2, size=8, color=AMBER)

    fp = sec.footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = fp.add_run(
        "Sources: QA.2B.1 (Second–Third) + QA.3.0 (Fourth–Sixth)  ·  Vision: P2.5.8 shadow only  ·  Confidential"
    )
    _set_run(fr, size=7.5, color=STEEL)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    _set_run(p.add_run("STEEL BEAM ESTIMATION"), size=20, bold=True, color=NAVY)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    _set_run(p.add_run("Current Performance & Vision Impact  ·  Second to Sixth Sets"), size=12, color=STEEL)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    _set_run(
        p.add_run(
            "Deterministic model = production authority.  Claude Vision = shadow / research only.  "
            "First Set is excluded from all aggregate KPIs."
        ),
        size=9,
        color=AMBER,
        bold=True,
    )

    heading_bar(doc, "CURRENT DETERMINISTIC MODEL  —  SECOND TO SIXTH SETS")
    kpis(
        doc,
        [
            ("Beam identification", "85.12%"),
            ("Bar identification", "50.37%"),
            ("Correct of detected bars", "33.95%"),
            ("Diameter identification", "75.67%"),
        ],
    )
    kpis(
        doc,
        [
            ("Bar matching (vs GT)", "17.10%"),
            ("Steel accuracy (kg-pooled)", "73.59%"),
            ("Overall accuracy", "60.76%"),
            ("Vision (Fifth Set only)", "+3.02 pp shadow"),
        ],
    )

    body(
        doc,
        "The model identifies 515 of 605 ground-truth beams (85.12%) and 2,100 of 4,169 "
        "ground-truth bar lines (50.37%). Of the 2,100 bars it does identify, 713 are a full "
        "MATCH on role, diameter and quantity (33.95%). Diameter is correct on 75.67% of those "
        "detected bars (WRONG_DIAMETER excluded). Combined steel is 109,580 kg versus "
        "148,911 kg estimator (73.59% QA.2A steel accuracy on pooled kg). Overall accuracy "
        "is the established four-metric mean: beam identification, bar identification, "
        "correct identification of detected bars, and pooled steel accuracy.",
        size=10,
        space_after=6,
    )
    body(
        doc,
        "These figures are the current production-authority result. They do not include Claude Vision.",
        size=10,
        bold=True,
        space_after=8,
    )

    heading_bar(doc, "1.  Drawing-wise deterministic performance")
    grid(
        doc,
        [
            "Drawing set",
            "Beam ID",
            "Bar ID",
            "Correct of detected",
            "Diameter ID",
            "Bar matching",
            "Steel",
            "Overall",
        ],
        [
            ["Second Set (known)", "95.52%", "71.76%", "35.74%", "77.98%", "25.65%", "85.50%", "72.13%"],
            ["Third Set (known)", "90.48%", "63.52%", "17.91%", "68.24%", "11.37%", "87.89%", "64.95%"],
            ["Fourth Set (unseen)", "78.32%", "39.89%", "37.33%", "73.07%", "14.89%", "58.39%", "53.48%"],
            ["Fifth Set (unseen)", "76.47%", "40.25%", "34.30%", "74.23%", "13.80%", "61.69%", "53.18%"],
            ["Sixth Set (unseen)", "95.86%", "61.45%", "38.87%", "81.63%", "23.89%", "96.54%", "73.18%"],
            ["ALL FIVE SETS", "85.12%", "50.37%", "33.95%", "75.67%", "17.10%", "73.59%", "60.76%"],
        ],
        [3.2, 1.7, 1.7, 2.2, 1.8, 1.9, 1.7, 1.8],
        total_last=True,
    )
    body(
        doc,
        "Beam ID = detected beams / GT beams.  Bar ID = detected bar lines / GT bar lines.  "
        "Correct of detected = MATCH / detected (QA.2A bar_accuracy).  "
        "Bar matching = MATCH / GT bars.  Diameter ID = (detected − WRONG_DIAMETER) / detected.  "
        "Steel = QA.2A 100 − |model−estimator| / estimator; five-set row is kg-pooled, not an average of %.  "
        "Overall = mean of Beam ID, Bar ID, Correct of detected, and Steel.  "
        "Known = Second–Third (development).  Unseen = Fourth–Sixth (generalization).",
        size=8,
        color=STEEL,
        space_after=8,
    )

    heading_bar(doc, "2.  What the five questions mean in numbers")
    grid(
        doc,
        ["Question", "KPI", "Five-set result", "Raw counts"],
        [
            ["What % of beams does the model identify?", "Beam identification", "85.12%", "515 / 605 GT beams"],
            ["What % of bars does the model identify?", "Bar identification", "50.37%", "2,100 / 4,169 GT bar lines"],
            ["Of bars identified, what % are correct?", "Correct of detected bars", "33.95%", "713 MATCH / 2,100 detected"],
            ["How accurately are diameters identified?", "Diameter identification", "75.67%", "(2,100 − 511 WRONG_DIA) / 2,100"],
            ["How close is steel quantity to the estimator?", "Steel accuracy", "73.59%", "109,580 / 148,911 kg"],
        ],
        [5.4, 3.4, 2.4, 5.8],
    )

    heading_bar(doc, "3.  Known (Second–Third) vs unseen (Fourth–Sixth)")
    grid(
        doc,
        ["Metric", "Second–Third known", "Fourth–Sixth unseen", "All five sets"],
        [
            ["Beam identification", "93.08%  (121/130)", "82.95%  (394/475)", "85.12%"],
            ["Bar identification", "67.25%  (573/852)", "46.04%  (1,527/3,317)", "50.37%"],
            ["Correct of detected bars", "26.53%  (152/573)", "36.74%  (561/1,527)", "33.95%"],
            ["Diameter identification", "73.00%", "76.69%", "75.67%"],
            ["Bar matching vs GT", "17.84%  (152/852)", "16.91%  (561/3,317)", "17.10%"],
            ["Steel accuracy (QA.2A set-mean)", "86.70%", "72.21%", "78.00%*"],
            ["Steel accuracy (kg-pooled)", "98.37%†", "68.23%", "73.59%"],
            ["Overall (pooled steel in five-set)", "68.39%‡", "59.49%‡", "60.76%"],
        ],
        [4.6, 4.2, 4.4, 3.8],
    )
    body(
        doc,
        "* QA.2A dashboard steel for multiple sets is the unweighted mean of per-set steel %.  "
        "† Known kg-pooled steel is inflated by Second-Set over-estimate cancelling Third-Set under-estimate — "
        "do not read 98% as known-set quality.  ‡ Group overall uses QA.2A set-mean steel.  "
        "Five-set overall uses kg-pooled steel.",
        size=8,
        color=STEEL,
        space_after=4,
    )

    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)

    heading_bar(doc, "4.  Diameter identification (detected bar lines)")
    body(
        doc,
        "QA.2A’s published field diameter_accuracy_pct is an alias of bar matching accuracy "
        "(MATCH / detected) and is not used here. Diameter identification below is counted from "
        "Bar Matching rows: a detected bar is diameter-correct unless status is WRONG_DIAMETER. "
        "GT diameter is the estimator line diameter. Ø32 has only 8 detected lines — percentage unstable.",
        size=9,
        space_after=5,
    )
    grid(
        doc,
        ["Diameter", "GT bar lines", "Detected", "MATCH", "WRONG_DIA", "Diameter ID", "Note"],
        [
            ["Ø8", "358", "159", "37", "3", "98.11%", "Usually right when found; many missing"],
            ["Ø10", "714", "235", "72", "10", "95.74%", "Same pattern — detection is the gap"],
            ["Ø12", "573", "224", "64", "31", "86.16%", ""],
            ["Ø16", "588", "404", "155", "140", "65.35%", "Frequent diameter swaps"],
            ["Ø20", "784", "431", "151", "186", "56.84%", "Weakest major diameter"],
            ["Ø25", "1,097", "639", "231", "136", "78.72%", "Includes many spacer lines"],
            ["Ø32", "26", "8", "3", "5", "37.50%", "Low volume — unstable"],
            ["TOTAL scored Ø8–Ø32", "4,140", "2,100", "713", "511", "75.67%", "Five-set headline"],
        ],
        [2.4, 2.2, 1.8, 1.6, 2.0, 2.2, 4.8],
        total_last=True,
    )

    heading_bar(doc, "5.  Diameter-wise steel quantity  —  not the same as diameter identification")
    body(
        doc,
        "Quantity ratio = automated kg / estimated kg × 100. It is not accuracy. "
        "A ratio above 100% is an overestimate. Ø16 is 143% of estimator kg while Ø10 is 50%.",
        size=9,
        space_after=5,
    )
    grid(
        doc,
        ["Diameter", "Estimated kg", "Automated kg", "Difference kg", "Abs % diff", "Quantity ratio"],
        [
            ["Ø8", "6,515", "4,389", "−2,126", "32.6%", "67%"],
            ["Ø10", "24,555", "12,236", "−12,319", "50.2%", "50%"],
            ["Ø12", "22,253", "11,670", "−10,582", "47.6%", "52%"],
            ["Ø16", "11,598", "16,564", "+4,965", "42.8%", "143%"],
            ["Ø20", "35,888", "29,320", "−6,568", "18.3%", "82%"],
            ["Ø25", "45,294", "32,965", "−12,328", "27.2%", "73%"],
            ["Ø32", "2,808", "2,436", "−372", "13.2%", "87%"],
            ["TOTAL", "148,911", "109,580", "−39,330", "26.4%", "74%"],
        ],
        [2.4, 3.0, 3.0, 2.8, 2.4, 3.4],
        total_last=True,
    )

    heading_bar(doc, "6.  Error profile  —  Second to Sixth Sets")
    grid(
        doc,
        ["Error category", "Second", "Third", "Fourth", "Fifth", "Sixth", "Five-set"],
        [
            ["Missing bars", "109", "170", "565", "870", "355", "2,069"],
            ["Wrong quantity", "81", "114", "103", "166", "176", "640"],
            ["Wrong diameter", "61", "94", "101", "151", "104", "511"],
            ["Extra bars", "74", "94", "51", "123", "81", "423"],
            ["Wrong role", "30", "23", "20", "45", "40", "158"],
            ["Beam missing", "3", "6", "31", "44", "6", "90"],
            ["Beam mismatch", "0", "4", "6", "0", "4", "14"],
            ["Steel difference", "1", "1", "1", "1", "1", "5"],
            ["Total classified errors", "359", "506", "878", "1,400", "767", "3,910"],
        ],
        [3.4, 1.9, 1.9, 1.9, 1.9, 1.9, 2.1],
        total_last=True,
    )
    body(
        doc,
        "Missing bars dominate unseen sets. Diameter and quantity errors remain material even when a bar is detected. "
        "This is why 50% bar identification does not become 50% correct steel.",
        size=9,
        space_after=8,
    )

    heading_bar(doc, "7.  Claude Vision  —  experimental, Fifth Set only")
    body(
        doc,
        "Engineering-level Vision-assisted recompute has currently been demonstrated on the Fifth Set. "
        "No valid engineering recompute exists for Second, Third, Fourth or Sixth Sets. "
        "A five-set Vision KPI is therefore not calculated. P2.5.8 is the authoritative Vision engineering result. "
        "P2.5.9–P2.5.11 are later safety/arbitration experiments and are not production Vision accuracy.",
        size=10,
        space_after=6,
    )
    grid(
        doc,
        ["Fifth Set steel", "Deterministic (production)", "Vision-assisted (P2.5.8 shadow)", "Estimator"],
        [
            ["Steel kg", "36,271.794", "38,049.747", "58,796.332"],
            ["Steel accuracy", "61.69%", "64.71%", "—"],
            ["Absolute error", "38.31%", "35.29%", "—"],
        ],
        [4.2, 4.4, 5.0, 3.4],
    )
    kpis(
        doc,
        [
            ("Accuracy lift", "+3.02 pp"),
            ("Error reduction", "7.88%"),
            ("Stirrup accuracy", "35.44% → 44.35%"),
            ("Beams  ↑ / = / ↓", "21 / 150 / 14"),
        ],
    )
    body(
        doc,
        "P2.5.8 replayed frozen P2.5.7 Claude results (0 new API calls) onto a sandbox copy of the "
        "reinforcement model. Practical value was OCR-difficult stirrups and slash spacing such as "
        "100/125/100 that SI.1 truncated to the first token. Promoted fields: diameter 29, legs 29, "
        "spacing 35, role 0. Production mutation = 0; steel / BBS / Excel production difference = 0.",
        size=9.5,
        space_after=5,
    )
    table = doc.add_table(rows=2, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell_text(
        table.cell(0, 0),
        "P2.5.8 DECISION  ·  NEGATIVE  ·  NOT PRODUCTION-SAFE",
        bold=True,
        color=WHITE,
        size=11,
        fill=RED_FILL,
    )
    cell_text(
        table.cell(1, 0),
        "Do not present 64.71% as current production accuracy. Fourteen beams worsened. "
        "Unrestricted new-stirrup inserts were rejected. P2.5.9–P2.5.11 gate those inserts "
        "(UNKNOWN-only, then evidence HOLD) and drop the lift to about +0.3 pp with zero worsenings.",
        color=RED,
        size=9,
        fill="F8E8E8",
    )
    set_widths(table, [17.0])
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(8)

    heading_bar(doc, "8.  Current workflow value  —  model-assisted, not autonomous")
    body(
        doc,
        "DXF  →  beam discovery  →  reinforcement interpretation  →  stirrup interpretation  →  "
        "steel calculation  →  BBS  →  Excel output  →  estimator verification.",
        size=10,
        bold=True,
        space_after=4,
    )
    body(
        doc,
        "Already automated: beam discovery (~85% of GT beams), annotation parsing, stirrup parsing, "
        "IS 456 steel calculation, BBS and Excel. Still required: finding the ~15% missed beams, "
        "the ~50% missed bars, and correcting ~66% of detected bars that are not a full MATCH. "
        "Positioning: model-assisted estimation / estimator first-pass support.",
        size=9.5,
        space_after=8,
    )

    heading_bar(doc, "9.  Indicative estimator time-saving potential")
    body(
        doc,
        "No estimator time-and-motion study, manual-vs-model hours, or productivity log exists in the project. "
        "Accuracy percentages are not time saved. 50% bar identification is not 50% time saved, because "
        "missed bars and incorrect detected bars still require drawing reading.",
        size=10,
        space_after=5,
    )
    grid(
        doc,
        ["If used today as a first-pass worksheet", "Indicative range", "Basis (not a measurement)"],
        [
            ["Beam take-off / marking", "Partial save", "85% of GT beams already listed"],
            ["Bar schedule compilation", "Limited save", "Half of GT bars missing; 2 in 3 detected bars need edit"],
            ["Steel calc / BBS / Excel typing", "High save", "Fully automated for whatever the model produced"],
            ["Verification / red-line", "Still required", "Unseen-set steel 58–62% except Sixth Set"],
            ["Overall estimation time", "20–35%", "Indicative only — requires time-and-motion study"],
        ],
        [5.4, 2.8, 8.8],
    )
    body(
        doc,
        "Indicative time-saving potential today: 20–35% of manual estimation time, if the Excel is used as a "
        "first pass and the estimator still verifies every beam. This is not reliably measurable until a "
        "timed estimator study is run. Computer pipeline runtime (e.g. Fifth Set ~2.8 h) is model processing, not estimator saving.",
        size=9.5,
        space_after=8,
    )

    heading_bar(doc, "10.  Current status")
    grid(
        doc,
        ["Item", "Value"],
        [
            ["Beam identification", "85.12%"],
            ["Bar identification", "50.37%"],
            ["Correct identification of detected bars", "33.95%"],
            ["Diameter identification accuracy", "75.67%"],
            ["Bar matching accuracy (MATCH / GT)", "17.10%"],
            ["Deterministic steel accuracy (kg-pooled)", "73.59%"],
            ["Overall accuracy", "60.76%"],
            ["Vision-assisted engineering improvement", "Fifth Set only: 61.69% → 64.71%  (+3.02 pp, shadow)"],
            ["Indicative estimator time saving", "20–35%  ·  not yet reliably measurable"],
            ["Current usage", "Model-assisted estimation / estimator first-pass support"],
            ["Production authority", "Deterministic model"],
            ["Claude Vision", "Controlled shadow / research  ·  P2.5.8 NEGATIVE for unrestricted promotion"],
        ],
        [6.5, 10.5],
    )
    body(
        doc,
        "Claude Vision has demonstrated a real engineering lift on difficult stirrup/OCR cases, "
        "but unrestricted promotion worsened 14 Fifth-Set beams. P2.5.9–P2.5.11 are making recovery "
        "engineering-safe rather than maximising aggregate accuracy. Production readiness is not claimed.",
        size=10,
        space_after=6,
    )

    heading_bar(doc, "11.  Methodology, sources, limitations")
    body(
        doc,
        "Second–Third: Version 9 QA.2B.1 GroundTruth_Benchmark_Report.xlsx (known development sets).  "
        "Fourth–Sixth: Version 10 QA.3.0 Generalization_Benchmark_Report.xlsx and per-set benchmark_result.json "
        "(unseen production runs qa2_*_20260806).  Counts are Excel Bar Detection / Bar Matching integers, "
        "not rounded percentages from the old six-set PDF.  Steel kg from QA.2A metric8.  "
        "Diameter kg from Diameter Steel Quantity sheets.  Vision from P2.5.8_STATUS.md.  "
        "First Set excluded (GT 18 beams / 85 bars / 1,424 kg).  "
        "Limitations: no estimator time study; diameter ID is status-based (PARTIAL_MATCH residual exists); "
        "Ø1/Ø2 rows in quantity sheets are parser artefacts and were omitted; "
        "Vision engineering recompute exists only for Fifth Set.",
        size=8.5,
        color=STEEL,
        space_after=2,
    )

    doc.save(DOCX)
    return DOCX


def export_pdf(docx_path: Path, pdf_path: Path) -> bool:
    try:
        import win32com.client  # type: ignore
    except ImportError:
        return False
    word = win32com.client.Dispatch("Word.Application")
    word.Visible = False
    try:
        doc = word.Documents.Open(str(docx_path))
        # 17 = wdFormatPDF
        doc.SaveAs(str(pdf_path), FileFormat=17)
        doc.Close(False)
        return pdf_path.exists()
    finally:
        word.Quit()


if __name__ == "__main__":
    path = build()
    print("DOCX", path)
    try:
        ok = export_pdf(path.resolve(), PDF.resolve())
        print("PDF", PDF if ok else "conversion_failed")
    except Exception as exc:
        print("PDF_ERROR", type(exc).__name__, exc)
