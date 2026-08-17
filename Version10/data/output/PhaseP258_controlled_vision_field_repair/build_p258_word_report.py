"""One-shot builder for the P2.5.8 Word status report."""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import nsmap, qn
from docx.shared import Cm, Emu, Pt, RGBColor

NAVY = RGBColor(0x1B, 0x36, 0x5D)
STEEL = RGBColor(0x2F, 0x5D, 0x7C)
GOLD = RGBColor(0xB8, 0x8A, 0x2E)
GREEN = RGBColor(0x1F, 0x6B, 0x3A)
RED = RGBColor(0x9B, 0x2C, 0x2C)
AMBER = RGBColor(0x8A, 0x5A, 0x00)
SLATE = RGBColor(0x3D, 0x4A, 0x57)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
ROW_ALT = "F4F7FA"
HEADER_FILL = "1B365D"
ACCENT_FILL = "2F5D7C"
GOLD_FILL = "B88A2E"
GREEN_FILL = "1F6B3A"
RED_FILL = "9B2C2C"
AMBER_FILL = "F4E6C4"


def _set_run(run, *, size=11, bold=False, color=SLATE, font="Calibri"):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = color


def _shade_cell(cell, hex_color: str) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def _set_cell_borders(cell, color="C5CDD6") -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        tcBorders.append(el)
    tcPr.append(tcBorders)


def _set_cell_text(cell, text, *, bold=False, color=SLATE, size=10.5, align="left", fill=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
    }[align]
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.08
    run = p.add_run(str(text))
    _set_run(run, size=size, bold=bold, color=color)
    if fill:
        _shade_cell(cell, fill)
    _set_cell_borders(cell)
    # vertical-ish padding via cell margins
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for edge, val in (("left", "80"), ("right", "80"), ("top", "40"), ("bottom", "40")):
        node = OxmlElement(f"w:{edge}")
        node.set(qn("w:w"), val)
        node.set(qn("w:type"), "dxa")
        tcMar.append(node)
    tcPr.append(tcMar)


def _set_table_widths(table, widths_cm):
    table.autofit = False
    table.allow_autofit = False
    total = sum(widths_cm)
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
    tblW = tblPr.find(qn("w:tblW"))
    if tblW is None:
        tblW = OxmlElement("w:tblW")
        tblPr.append(tblW)
    tblW.set(qn("w:w"), str(int(total * 567)))
    tblW.set(qn("w:type"), "dxa")
    grid = tbl.find(qn("w:tblGrid"))
    if grid is not None:
        for i, w in enumerate(widths_cm):
            grid_col = list(grid)[i]
            grid_col.set(qn("w:w"), str(int(w * 567)))
    for row in table.rows:
        for i, w in enumerate(widths_cm):
            row.cells[i].width = Cm(w)


def add_heading_bar(doc, title: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    _set_cell_text(cell, title, bold=True, color=WHITE, size=12, fill=HEADER_FILL)
    _set_table_widths(table, [16.5])
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def add_kv_table(doc, rows, col_widths=(5.5, 11.0)):
    table = doc.add_table(rows=len(rows), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (k, v) in enumerate(rows):
        fill = HEADER_FILL if i == 0 and k == "__header__" else (ROW_ALT if i % 2 else "FFFFFF")
        if k == "__header__":
            continue
        _set_cell_text(table.cell(i, 0), k, bold=True, color=NAVY, size=10.5, fill=fill)
        _set_cell_text(table.cell(i, 1), v, color=SLATE, size=10.5, fill=fill)
    _set_table_widths(table, list(col_widths))
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(10)


def add_grid_table(doc, headers, data, col_widths, highlight_last=False):
    table = doc.add_table(rows=1 + len(data), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(headers):
        _set_cell_text(table.cell(0, j), h, bold=True, color=WHITE, size=10, align="center", fill=HEADER_FILL)
    for i, row in enumerate(data):
        fill = ROW_ALT if i % 2 else "FFFFFF"
        for j, val in enumerate(row):
            align = "left" if j == 0 else "center"
            color = SLATE
            bold = False
            if highlight_last and i == len(data) - 1:
                fill = "E8F1EA"
                color = GREEN
                bold = True
            _set_cell_text(
                table.cell(i + 1, j),
                val,
                bold=bold or j == 0,
                color=NAVY if j == 0 else color,
                size=10.5,
                align=align,
                fill=fill,
            )
    _set_table_widths(table, col_widths)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(10)


def add_kpi_row(doc, kpis):
    table = doc.add_table(rows=2, cols=len(kpis))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    fills = [HEADER_FILL, ACCENT_FILL, GREEN_FILL, GOLD_FILL]
    for i, (label, value) in enumerate(kpis):
        fill = fills[i % len(fills)]
        _set_cell_text(table.cell(0, i), label, bold=True, color=WHITE, size=9, align="center", fill=fill)
        _set_cell_text(table.cell(1, i), value, bold=True, color=NAVY, size=16, align="center", fill="FFFFFF")
    width = 16.5 / len(kpis)
    _set_table_widths(table, [width] * len(kpis))
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(12)


def add_body(doc, text, *, size=11, color=SLATE, space_after=8):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    _set_run(run, size=size, color=color)
    return p


def add_decision_banner(doc):
    table = doc.add_table(rows=2, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_cell_text(
        table.cell(0, 0),
        "FINAL DECISION  ·  NEGATIVE",
        bold=True,
        color=WHITE,
        size=13,
        align="center",
        fill=RED_FILL,
    )
    _set_cell_text(
        table.cell(1, 0),
        "Tighten validation / revert promotion class.  "
        "Accuracy improved, but 14 beams worsened — not production-safe.",
        bold=False,
        color=RED,
        size=11,
        align="center",
        fill="F8E8E8",
    )
    _set_table_widths(table, [16.5])
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(12)


def build() -> Path:
    out = Path(__file__).with_name("P2.5.8_Vision_Field_Repair_Report.docx")
    doc = Document()

    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    section.top_margin = Cm(1.6)
    section.bottom_margin = Cm(1.8)

    # Header
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = hp.add_run("SteelBeam Estimator  ·  Version 10  ·  Confidential research")
    _set_run(r, size=9, bold=True, color=STEEL)
    r2 = hp.add_run("                                          MODEL 10.8.4")
    _set_run(r2, size=9, color=GOLD)

    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = fp.add_run("P2.5.8 Controlled Vision Field-Repair  ·  Shadow evaluation only  ·  Production write = False")
    _set_run(fr, size=8, color=STEEL)

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("P2.5.8 STATUS REPORT")
    _set_run(r, size=22, bold=True, color=NAVY, font="Calibri")

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run("Controlled Vision Field-Repair & Engineering Recompute")
    _set_run(r, size=13, color=STEEL)

    add_kpi_row(
        doc,
        [
            ("Baseline accuracy", "61.69%"),
            ("Vision-assisted", "64.71%"),
            ("Improvement", "+3.02 pp"),
            ("Final decision", "NEGATIVE"),
        ],
    )

    add_body(
        doc,
        "This report records the shadow recompute of Fifth Set steel after overlaying frozen "
        "P2.5.7 Claude Vision interpretations onto a sandbox copy of the production reinforcement "
        "model. No new Claude API calls were made. Production steel, BBS, and Excel were not mutated.",
        size=11,
    )

    add_heading_bar(doc, "1.  Identity")
    add_kv_table(
        doc,
        [
            ("Model version", "10.8.4"),
            ("Phase ID", "P2.5.8"),
            ("Phase name", "Controlled Vision Field-Repair & Engineering Recompute"),
            ("Phase status", "PASS  (experiment completed)"),
            ("Final decision", "NEGATIVE  (not production-safe)"),
            ("Engineering changes", "NONE"),
            ("Vision mode", "REPLAY_P257_LIVE_RESULTS"),
        ],
    )

    add_heading_bar(doc, "2.  Dataset")
    add_kv_table(
        doc,
        [
            ("Drawing set", "Fifth Set Drawings"),
            ("DXF count", "3"),
            ("Unique model-detected beams", "143"),
            ("Annotation candidates", "401"),
        ],
    )

    add_heading_bar(doc, "3.  Vision intake")
    add_body(
        doc,
        "Claude results were replayed from the frozen P2.5.7 live capture (vision_results.json). "
        "The official P2.5.8 run spent $0 on new API calls.",
        size=10.5,
        space_after=6,
    )
    add_grid_table(
        doc,
        ["Metric", "Value"],
        [
            ["Candidates available", "41"],
            ["Candidates eligible for promotion", "35"],
            ["Fields promoted", "93"],
            ["Fields blocked", "30"],
            ["New Claude calls", "0"],
            ["Claude cost (USD)", "0.00"],
            ["Replay", "True"],
        ],
        [8.0, 8.5],
    )

    add_heading_bar(doc, "4.  Field impact")
    add_grid_table(
        doc,
        ["Field", "Before", "After", "Promoted"],
        [
            ["Diameter", "6", "35", "29"],
            ["Legs", "6", "35", "29"],
            ["Spacing", "0", "35", "35"],
            ["Reinforcement role", "35", "35", "0"],
        ],
        [5.5, 3.5, 3.5, 4.0],
    )

    add_heading_bar(doc, "5.  Engineering impact  —  headline result")
    add_body(
        doc,
        "Accuracy is model steel ÷ estimator steel on the Fifth Set. Vision overlay recovered "
        "incomplete stirrup schedules that the SI.1 parser truncated, lifting total accuracy by "
        "3.02 percentage points.",
        size=10.5,
        space_after=6,
    )
    add_grid_table(
        doc,
        ["Measure", "Baseline (deterministic)", "Vision-assisted", "Estimator"],
        [
            ["Steel (kg)", "36,271.794", "38,049.747", "58,796.332"],
            ["Accuracy", "61.69%", "64.71%", "—"],
            ["Absolute steel error", "38.31%", "35.29%", "—"],
        ],
        [4.2, 4.3, 4.0, 4.0],
        highlight_last=False,
    )

    kpi2 = doc.add_table(rows=1, cols=2)
    kpi2.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_cell_text(
        kpi2.cell(0, 0),
        "Accuracy improvement   +3.02 percentage points",
        bold=True,
        color=WHITE,
        size=12,
        align="center",
        fill=GREEN_FILL,
    )
    _set_cell_text(
        kpi2.cell(0, 1),
        "Error reduction   7.88%",
        bold=True,
        color=WHITE,
        size=12,
        align="center",
        fill=ACCENT_FILL,
    )
    _set_table_widths(kpi2, [8.25, 8.25])
    doc.add_paragraph().paragraph_format.space_after = Pt(10)

    add_heading_bar(doc, "6.  Stirrup impact")
    add_body(
        doc,
        "Almost all of the steel lift is stirrup recovery. Vision restored slash spacing patterns "
        "(for example 100/125/100) that the deterministic parser captured only as the first token.",
        size=10.5,
        space_after=6,
    )
    add_grid_table(
        doc,
        ["Measure", "Before", "After", "Estimator"],
        [
            ["Stirrup steel (kg)", "7,075.317", "8,853.270", "19,962.308"],
            ["Stirrup accuracy", "35.44%", "44.35%", "—"],
            ["Stirrup quantity (pcs)", "4,928", "6,027", "13,869"],
        ],
        [4.5, 4.0, 4.0, 4.0],
    )
    note = doc.add_paragraph()
    note.paragraph_format.space_after = Pt(10)
    r = note.add_run("Stirrup accuracy improvement:  +8.91 percentage points")
    _set_run(r, size=11, bold=True, color=GREEN)

    add_heading_bar(doc, "7.  Beam impact")
    add_grid_table(
        doc,
        ["Outcome", "Count", "Notes"],
        [
            ["Improved", "21", "Closer to estimator after Vision overlay"],
            ["Unchanged", "150", "No material kg change"],
            ["Worsened", "14", "Unsafe inserts — primary reason for NEGATIVE"],
            ["Newly resolved", "0", "—"],
            ["Still unresolved", "185", "Estimator ∪ model universe"],
        ],
        [4.0, 3.0, 9.5],
    )
    add_body(
        doc,
        "The 14 worsenings are why the experiment is not production-safe despite the +3.02 pp "
        "headline. Later phases (P2.5.9–P2.5.11) gate new-stirrup inserts to drive worsenings to zero, "
        "at the cost of most of this accuracy lift.",
        size=10.5,
    )

    add_heading_bar(doc, "8.  Safety firewall")
    add_grid_table(
        doc,
        ["Check", "Result"],
        [
            ["Conflicts", "12"],
            ["Blocked fields", "30"],
            ["Validation failures", "67"],
            ["Production mutations", "0"],
            ["Production output difference", "0"],
            ["Steel production difference", "0"],
            ["BBS production difference", "0"],
            ["Excel production difference", "0"],
        ],
        [8.0, 8.5],
    )

    add_heading_bar(doc, "9.  Regression, tests, and cost")
    add_grid_table(
        doc,
        ["Item", "Result"],
        [
            ["P2.5.1", "PASS"],
            ["P2.5.4", "PASS"],
            ["P2.5.5", "PASS"],
            ["P2.5.6", "PASS"],
            ["P2.5.7", "PASS"],
            ["Fingerprint unchanged", "True"],
            ["Unit tests", "24 / 24"],
            ["New Claude calls / tokens in / out", "0 / 0 / 0"],
            ["Estimated cost", "USD 0.00  (replay)"],
        ],
        [8.0, 8.5],
    )

    add_heading_bar(doc, "10.  Decision")
    add_decision_banner(doc)
    add_body(
        doc,
        "Recommendation: keep P2.5.8 RESEARCH_ONLY. Do not promote this overlay class into "
        "production. Subsequent phases must require stronger evidence before inserting a new stirrup "
        "onto a beam that previously had none.",
        size=11,
    )

    src = doc.add_paragraph()
    src.paragraph_format.space_before = Pt(8)
    r = src.add_run(
        "Source: Version10/data/output/PhaseP258_controlled_vision_field_repair/P2.5.8_STATUS.md"
    )
    _set_run(r, size=8, color=STEEL)

    doc.save(out)
    return out


if __name__ == "__main__":
    path = build()
    print(path)
