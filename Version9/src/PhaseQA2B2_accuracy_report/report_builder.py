"""
QA.2B.2 — Build Overall_Accuracy_Report_V9.6.1.docx from the V8.9.1 template.
MODEL_VERSION: 9.6.2
"""
from __future__ import annotations

import json
import shutil
from copy import deepcopy
from pathlib import Path
from typing import Any, Dict, List, Optional

from docx import Document
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor
from docx.text.paragraph import Paragraph

MODEL_VERSION = "9.6.2"
PHASE_ID = "QA.2B.2"

BASELINE = {
    "beam_detection_pct": 93.92,
    "bar_detection_pct": 41.52,
    "bar_accuracy_pct": 24.16,
    "steel_accuracy_pct": 72.69,
    "overall_accuracy_pct": 58.07,
}

RED = RGBColor(0xEE, 0x00, 0x00)
GREEN = RGBColor(0x00, 0x80, 0x00)


def _pct(v: float, digits: int = 2) -> str:
    return f"{float(v):.{digits}f}%"


def _delta(curr: float, base: float) -> str:
    d = float(curr) - float(base)
    sign = "+" if d >= 0 else ""
    return f"{sign}{d:.2f}%"


def _clear_paragraph(para: Paragraph) -> None:
    for run in list(para.runs):
        run._element.getparent().remove(run._element)


def _heading(para: Paragraph, text: str, *, size_pt: Optional[float] = None) -> None:
    _clear_paragraph(para)
    run = para.add_run(text)
    run.bold = True
    if size_pt is not None:
        run.font.size = Pt(size_pt)


def _label_value(para: Paragraph, label: str, value: str) -> None:
    _clear_paragraph(para)
    para.add_run(label)
    r = para.add_run(value)
    r.bold = True


def _metric_line(para: Paragraph, label: str, value: str) -> None:
    _clear_paragraph(para)
    r0 = para.add_run(f"{label} ")
    r0.bold = True
    r1 = para.add_run(value)
    r1.bold = True


def _plain(para: Paragraph, text: str) -> None:
    _clear_paragraph(para)
    para.add_run(text)


def _insert_paragraph_after(paragraph: Paragraph, text: str = "", *, bold: bool = False) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if text:
        run = new_para.add_run(text)
        run.bold = bold
    return new_para


def _insert_paragraph_before(paragraph: Paragraph, text: str = "", *, bold: bool = False) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addprevious(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if text:
        run = new_para.add_run(text)
        run.bold = bold
    return new_para


def _insert_table_after(paragraph: Paragraph, rows: int, cols: int):
    doc = paragraph.part.document
    table = doc.add_table(rows=rows, cols=cols)
    tbl = table._tbl
    paragraph._p.addnext(tbl)
    return table


def _set_cell(cell, text: str, *, bold: bool = False, color: Optional[RGBColor] = None) -> None:
    para = cell.paragraphs[0]
    _clear_paragraph(para)
    run = para.add_run(text)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def _find_para(doc: Document, exact: str = "", *, startswith: str = "") -> Paragraph:
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if exact and t == exact:
            return p
        if startswith and t.startswith(startswith):
            return p
    raise KeyError(exact or startswith)


def _find_para_contains(doc: Document, needle: str) -> Paragraph:
    for p in doc.paragraphs:
        if needle in (p.text or ""):
            return p
    raise KeyError(needle)


def build_report(
    *,
    template_path: Path,
    qa2b1_dir: Path,
    output_docx: Path,
) -> Dict[str, Any]:
    template_path = Path(template_path)
    qa2b1_dir = Path(qa2b1_dir)
    output_docx = Path(output_docx)
    output_docx.parent.mkdir(parents=True, exist_ok=True)

    data = json.loads(
        (qa2b1_dir / "GroundTruth_Benchmark_Report.json").read_text(encoding="utf-8")
    )
    bench = data["benchmark"]
    errors = (data.get("errors") or {}).get("frequency") or {}
    sets = {d["drawing_set"]: d for d in data.get("drawing_sets") or []}

    shutil.copy2(template_path, output_docx)
    doc = Document(str(output_docx))

    curr = {
        "beam_detection_pct": float(bench["beam_detection_pct"]),
        "bar_detection_pct": float(bench["bar_detection_pct"]),
        "bar_accuracy_pct": float(bench["bar_accuracy_pct"]),
        "steel_accuracy_pct": float(bench["steel_accuracy_pct"]),
        "overall_accuracy_pct": float(bench["overall_accuracy_pct"]),
    }

    # 1) Title
    _heading(
        doc.paragraphs[0],
        "Steel Beam Estimation – Ground Truth Benchmark Summary",
        size_pt=14,
    )
    _heading(
        doc.paragraphs[1],
        "Version 9.6.1 | QA.2B.1 – Production Regeneration & Ground Truth Comparison",
    )

    # 2) Benchmark Results
    _heading(doc.paragraphs[2], "Benchmark Results")
    _plain(
        doc.paragraphs[3],
        "Production estimation outputs were regenerated from DXF for three independent beam drawing sets "
        "using the latest Version 9 engineering pipeline. No prior Estimation_Output.xlsx workbook was reused. "
        "For each drawing set, the regenerated production workbook was compared against the estimator-generated "
        "bar bending schedule (ground truth).",
    )

    # 3) Accuracy Improvement Summary (insert before Drawing-wise Performance)
    dw = _find_para(doc, "Drawing-wise Performance")
    h_imp = _insert_paragraph_before(dw, "Accuracy Improvement Summary", bold=True)
    intro = _insert_paragraph_after(
        h_imp,
        "The following table compares the Version 8.9.1 baseline benchmark against the Version 9.6.1 "
        "results from QA.2B.1 regenerated production outputs.",
    )
    imp_rows = [
        ("Metric", "Version 8.9.1", "Version 9.6.1", "Improvement"),
        (
            "Beam Detection",
            _pct(BASELINE["beam_detection_pct"]),
            _pct(curr["beam_detection_pct"]),
            _delta(curr["beam_detection_pct"], BASELINE["beam_detection_pct"]),
        ),
        (
            "Bar Detection",
            _pct(BASELINE["bar_detection_pct"]),
            _pct(curr["bar_detection_pct"]),
            _delta(curr["bar_detection_pct"], BASELINE["bar_detection_pct"]),
        ),
        (
            "Bar Matching Accuracy",
            _pct(BASELINE["bar_accuracy_pct"]),
            _pct(curr["bar_accuracy_pct"]),
            _delta(curr["bar_accuracy_pct"], BASELINE["bar_accuracy_pct"]),
        ),
        (
            "Steel Accuracy",
            _pct(BASELINE["steel_accuracy_pct"]),
            _pct(curr["steel_accuracy_pct"]),
            _delta(curr["steel_accuracy_pct"], BASELINE["steel_accuracy_pct"]),
        ),
        (
            "Overall Accuracy",
            _pct(BASELINE["overall_accuracy_pct"]),
            _pct(curr["overall_accuracy_pct"]),
            _delta(curr["overall_accuracy_pct"], BASELINE["overall_accuracy_pct"]),
        ),
    ]
    imp_table = _insert_table_after(intro, len(imp_rows), 4)
    try:
        # Prefer style of original drawing-wise table (still tables[0] before move? after insert may differ)
        for t in doc.tables:
            if t.rows and "Drawing Set" in (t.rows[0].cells[0].text or ""):
                imp_table.style = t.style
                break
    except Exception:
        pass
    for ri, row_vals in enumerate(imp_rows):
        for ci, val in enumerate(row_vals):
            color = None
            if ri > 0 and ci == 3:
                if val.startswith("+") and not val.startswith("+0.00"):
                    color = GREEN
                elif val.startswith("-"):
                    color = RED
            _set_cell(
                imp_table.rows[ri].cells[ci],
                val,
                bold=(ri == 0 or ci == 0),
                color=color,
            )
    _insert_paragraph_after(Paragraph(imp_table._tbl, intro._parent), "")

    # 4) Drawing-wise Performance table
    _heading(_find_para(doc, "Drawing-wise Performance"), "Drawing-wise Performance")
    t0 = None
    for t in doc.tables:
        if t.rows and "Drawing Set" in (t.rows[0].cells[0].text or ""):
            t0 = t
            break
    assert t0 is not None
    if len(t0.columns) == 4:
        for row in t0.rows:
            row._tr.append(deepcopy(row.cells[-1]._tc))

    def beam_cell(ds: Dict[str, Any]) -> str:
        bm = ds["beam_matching"]
        return f"{bm['detected_beams']} / {bm['estimator_beams']} ({bm['detection_pct']:.1f}%)"

    first, second, third = (
        sets["First Set Drawings"],
        sets["Second Set Drawings"],
        sets["Third Set Drawings"],
    )
    draw_rows = [
        ["Drawing Set", "Beam Detection", "Bar Detection", "Bar Accuracy", "Steel Accuracy"],
        [
            "First Set",
            beam_cell(first),
            _pct(first["bar_matching"]["detection_pct"]),
            _pct(first["bar_matching"]["accuracy_pct"]),
            _pct(first["steel"]["accuracy_pct"]),
        ],
        [
            "Second Set",
            beam_cell(second),
            _pct(second["bar_matching"]["detection_pct"]),
            _pct(second["bar_matching"]["accuracy_pct"]),
            _pct(second["steel"]["accuracy_pct"]),
        ],
        [
            "Third Set",
            beam_cell(third),
            _pct(third["bar_matching"]["detection_pct"]),
            _pct(third["bar_matching"]["accuracy_pct"]),
            _pct(third["steel"]["accuracy_pct"]),
        ],
        [
            "Overall",
            f"{bench['detected_beams']} / {bench['total_beams']} ({bench['beam_detection_pct']:.2f}%)",
            _pct(bench["bar_detection_pct"]),
            _pct(bench["bar_accuracy_pct"]),
            _pct(bench["steel_accuracy_pct"]),
        ],
    ]
    for ri, vals in enumerate(draw_rows):
        for ci, val in enumerate(vals):
            color = None
            if ri > 0 and ci == 3:
                try:
                    if float(val.rstrip("%")) < 50:
                        color = RED
                except ValueError:
                    pass
            _set_cell(
                t0.rows[ri].cells[ci],
                val,
                bold=(ri == 0 or ri == 4),
                color=color,
            )

    try:
        narr = _find_para_contains(doc, "consistent beam detection")
        _plain(
            narr,
            "The results indicate consistent beam detection across all drawing sets. "
            "Steel quantity accuracy has improved materially versus Version 8.9.1, while "
            "reinforcement bar interpretation remains the primary factor limiting overall estimation accuracy, "
            "especially on the Third drawing set.",
        )
    except KeyError:
        pass

    # 5) Overall Performance Summary
    _heading(_find_para(doc, "Overall Performance Summary"), "Overall Performance Summary")
    _heading(_find_para(doc, "Beam Detection"), "Beam Detection")
    _label_value(
        _find_para(doc, startswith="Ground Truth Beams:"),
        "Ground Truth Beams: ",
        str(bench["total_beams"]),
    )
    _label_value(
        _find_para(doc, startswith="Correctly Detected"),
        "Correctly Detected Beams: ",
        str(bench["correct_beams"]),
    )
    _metric_line(
        _find_para(doc, startswith="Beam Detection Accuracy"),
        "Beam Detection Accuracy:",
        _pct(bench["beam_detection_pct"]),
    )
    _heading(
        _find_para(doc, "Reinforcement Bar Interpretation"),
        "Reinforcement Bar Interpretation",
    )
    _label_value(
        _find_para(doc, startswith="Ground Truth Bars:"),
        "Ground Truth Bars: ",
        str(bench["total_bars"]),
    )
    _label_value(
        _find_para(doc, startswith="Detected Bars:"),
        "Detected Bars: ",
        str(bench["detected_bars"]),
    )
    _label_value(
        _find_para(doc, startswith="Correctly Matched Bars:"),
        "Correctly Matched Bars: ",
        str(bench["correct_bars"]),
    )
    _label_value(
        _find_para(doc, startswith="Missing Bars:"),
        "Missing Bars: ",
        str(bench["missing_bars"]),
    )
    _metric_line(
        _find_para(doc, startswith="Bar Detection Accuracy"),
        "Bar Detection Accuracy:",
        _pct(bench["bar_detection_pct"]),
    )
    _metric_line(
        _find_para(doc, startswith="Bar Matching Accuracy"),
        "Bar Matching Accuracy:",
        _pct(bench["bar_accuracy_pct"]),
    )
    _heading(_find_para(doc, "Steel Quantity Estimation"), "Steel Quantity Estimation")
    _plain(
        _find_para(doc, startswith="Average Steel Quantity Accuracy"),
        f"Average Steel Quantity Accuracy: {_pct(bench['steel_accuracy_pct'])}",
    )
    _heading(_find_para(doc, "Overall Model Accuracy"), "Overall Model Accuracy")
    _plain(
        _find_para(doc, startswith="Overall Ground Truth Benchmark Accuracy"),
        f"Overall Ground Truth Benchmark Accuracy: {_pct(bench['overall_accuracy_pct'])}",
    )

    # 6) Key Error Categories
    _heading(_find_para(doc, "Key Error Categories"), "Key Error Categories")
    try:
        _plain(
            _find_para_contains(doc, "automatically classified discrepancies"),
            "The benchmark automatically classified discrepancies between the regenerated "
            "model output and the estimator's ground truth.",
        )
    except KeyError:
        pass

    err_table = None
    for t in doc.tables:
        if t.rows and "Error" in (t.rows[0].cells[0].text or ""):
            err_table = t
            break
    assert err_table is not None
    err_rows = [
        ("Error Category", "Count"),
        ("Missing Bars", str(errors.get("Missing Bar", 0))),
        ("Wrong Quantity", str(errors.get("Wrong Quantity", 0))),
        ("Extra Bars", str(errors.get("Extra Bar", 0))),
        ("Wrong Diameter", str(errors.get("Wrong Diameter", 0))),
        ("Wrong Role Classification", str(errors.get("Wrong Role", 0))),
        ("Beam Missing", str(errors.get("Beam Missing", 0))),
        ("Beam Mismatch", str(errors.get("Beam Mismatch", 0))),
        ("Steel Difference", str(errors.get("Steel Difference", 0))),
    ]
    while len(err_table.rows) < len(err_rows):
        err_table.add_row()
    for ri, vals in enumerate(err_rows):
        _set_cell(err_table.rows[ri].cells[0], vals[0], bold=(ri == 0))
        _set_cell(err_table.rows[ri].cells[1], vals[1], bold=(ri == 0))
    for ri in range(len(err_rows), len(err_table.rows)):
        _set_cell(err_table.rows[ri].cells[0], "")
        _set_cell(err_table.rows[ri].cells[1], "")

    try:
        p = _find_para_contains(doc, "reinforcement interpretation")
        if "Most of" in p.text or "errors are associated" in p.text:
            _clear_paragraph(p)
            p.add_run("Most of the errors remain associated with ")
            rb = p.add_run("reinforcement interpretation")
            rb.bold = True
            p.add_run(
                ", particularly missing bars, wrong quantity, extra bars, and diameter resolution. "
                f"Total classified errors: {sum(int(v) for v in errors.values())}."
            )
    except KeyError:
        pass

    # 7) Key Engineering Improvements (before Key Findings)
    kf = _find_para(doc, "Key Findings")
    eng_h = _insert_paragraph_before(
        kf, "Key Engineering Improvements Since Version 8.9.1", bold=True
    )
    eng_intro = _insert_paragraph_after(
        eng_h,
        "Since the Version 8.9.1 baseline, the Version 9 engineering programme has delivered "
        "the following major capabilities that are now included in the regenerated benchmark:",
    )
    prev = eng_intro
    for item in [
        "Stirrup Recovery Engine",
        "OpenCV geometric evidence pipeline",
        "Beam Ownership Engine",
        "Annotation Graph Resolver",
        "Adaptive Render Extent",
        "Shared Engineering Ownership",
        "Shared Scope Deduplication",
        "End-to-End Pipeline Integration (QA.2B.0)",
        "Production Output Regeneration (QA.2B.1)",
    ]:
        prev = _insert_paragraph_after(prev, f"• {item}")
    _insert_paragraph_after(prev, "")

    # 8) Key Findings
    kf = _find_para(doc, "Key Findings")
    _heading(kf, "Key Findings")
    findings = [
        f"Beam detection remains highly reliable at {_pct(bench['beam_detection_pct'])}, "
        "confirming that beam discovery is largely stable across all three drawing sets.",
        f"Steel estimation accuracy has improved significantly to {_pct(bench['steel_accuracy_pct'])} "
        f"(from {_pct(BASELINE['steel_accuracy_pct'])} in Version 8.9.1).",
        f"Reinforcement interpretation has improved substantially: bar detection rose to "
        f"{_pct(bench['bar_detection_pct'])} (from {_pct(BASELINE['bar_detection_pct'])}), "
        f"and bar matching accuracy to {_pct(bench['bar_accuracy_pct'])} "
        f"(from {_pct(BASELINE['bar_accuracy_pct'])}).",
        "The Third drawing set remains the primary benchmark challenge, with the lowest "
        "bar matching accuracy among the three sets.",
        "Remaining work should focus on reinforcement interpretation, continuous bars, "
        "diameter resolution, and complex shared beam annotations.",
        f"Overall ground-truth benchmark accuracy is now {_pct(bench['overall_accuracy_pct'])}, "
        f"an improvement of {_delta(bench['overall_accuracy_pct'], BASELINE['overall_accuracy_pct'])} "
        "versus Version 8.9.1.",
    ]
    # Collect paragraphs after Key Findings until document end
    paras = list(doc.paragraphs)
    kf_i = next(i for i, p in enumerate(paras) if (p.text or "").strip() == "Key Findings")
    tail = paras[kf_i + 1 :]
    # Drop empty engineering leftovers shouldn't be after findings
    for i, text in enumerate(findings):
        if i < len(tail):
            _plain(tail[i], text)
        else:
            prev = tail[-1] if tail else kf
            np = _insert_paragraph_after(prev, text)
            tail.append(np)
    for p in tail[len(findings) :]:
        _clear_paragraph(p)

    doc.save(str(output_docx))

    meta = {
        "phase_id": PHASE_ID,
        "model_version": MODEL_VERSION,
        "template": str(template_path),
        "output": str(output_docx),
        "baseline": BASELINE,
        "current": curr,
        "improvement": {k: round(curr[k] - BASELINE[k], 2) for k in curr},
        "qa2b1_source": str(qa2b1_dir / "GroundTruth_Benchmark_Report.json"),
    }
    output_docx.with_suffix(".meta.json").write_text(
        json.dumps(meta, indent=2), encoding="utf-8"
    )
    return meta
